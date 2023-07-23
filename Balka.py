from PyQt5 import QtCore, QtGui, QtWidgets
import sys
import pickle
import Sortament

from math import pi
from math import sin
from okno_general import ui
from okno_general import Form
from okno_general import app

# from Tab73 import fsi73

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
document = Document()

# Рисунок
from PyQt5.QtWidgets import QWidget
from PyQt5.QtWidgets import QApplication
from PyQt5.QtGui import QPainter
from PyQt5.QtGui import QBrush
from PyQt5.QtGui import QPen
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap

_translate = QtCore.QCoreApplication.translate
Form.setWindowTitle(_translate("Form", "VXV"))
ui.tabWidget.setCurrentIndex(3)

# -----------------------------------------------------------------------------
# Вывод ошибки
def error_show(x):
    ui.label_3.setText(_translate("Form", "Ошибка"))
    ui.textEdit_3.setTextColor(QtGui.QColor (255, 0, 0))
    ui.textEdit_3.setText('Ошибка данных')
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_3.append(x)
    # QtWidgets.QMessageBox.information(Form, 'Ошибка данных', str(x))
# -----------------------------------------------------------------------------

# -----------------------------------------------------------------------------
def text_centr(x):
    ui.textEdit_3.append('')
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 100, 150)) # цвет текста
    # ui.textEdit_3.setFontItalic(True) # курсивный текст
    ui.textEdit_3.setFontWeight(100) # жирный текст
    ui.textEdit_3.append(x)
    ui.textEdit_3.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    ui.textEdit_3.setFontWeight(1) # убираем жирный текст
    # ui.textEdit_3.setFontItalic(False) # убираем курсивный текст
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 0, 0))
    global document
    # document.add_paragraph(x)
    paragraph = document.add_paragraph()
    # paragraph.add_run(x).bold = True
    paragraph.add_run(x, style='Intense Emphasis').bold = True
    # paragraph.add_run(x).font.color.rgb = RGBColor(0, 100, 150)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
def text_centr_black(x):
    ui.textEdit_3.append(x)
    ui.textEdit_3.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    global document
    paragraph = document.add_paragraph(x)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
def text_abzac(x):
    ui.textEdit_3.append('       {}'.format(x))
    ui.textEdit_3.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    global document
    document.add_paragraph(x)
def text_abzac_left(x):
    ui.textEdit_3.append('{}'.format(x))
    ui.textEdit_3.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    global document
    document.add_paragraph(x)
def text_abzac_color(x):
    ui.textEdit_3.append('')
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 100, 150))
    ui.textEdit_3.setFontWeight(100) # жирный текст
    ui.textEdit_3.append('       {}'.format(x))
    ui.textEdit_3.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_3.setFontWeight(1) # жирный текст
    global document
    paragraph = document.add_paragraph()
    paragraph.add_run(x, style='Intense Emphasis').bold = True

def text_centr_red(x):
    ui.textEdit_3.append('')
    ui.textEdit_3.setTextColor(QtGui.QColor (255, 0, 0)) # цвет текста
    # ui.textEdit_3.setFontItalic(True) # курсивный текст
    ui.textEdit_3.setFontWeight(100) # жирный текст
    ui.textEdit_3.append(x)
    ui.textEdit_3.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    ui.textEdit_3.setFontWeight(1) # убираем жирный текст
    # ui.textEdit_3.setFontItalic(False) # убираем курсивный текст
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 0, 0))
    global document
    # document.add_paragraph(x)
    paragraph = document.add_paragraph()
    # paragraph.add_run(x).bold = True
    paragraph.add_run(x, style='Intense Emphasis').bold = True
    # paragraph.add_run(x).font.color.rgb = RGBColor(0, 100, 150)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

def yvedomlenie(x):
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 100, 150))
    ui.textEdit_3.setFontWeight(100) # жирный текст
    ui.textEdit_3.append('{}'.format(x))
    ui.textEdit_3.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    ui.textEdit_3.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_3.setFontWeight(1) # жирный текст
# -----------------------------------------------------------------------------
# Проверка ячеек на пустые значения и ","
def vvod(nomerwidgeta, strok, stolbec):
    yoy = []
    for i in range(1, strok+1):
        x = eval('ui.tableWidget{}.item({}-1, {}).text()'.format(nomerwidgeta, i, stolbec))
        if x != '': 
            x = x.replace(',', '.')
        else: 
            x = 0
        if '.' in str(x):
            yoy.append(round(float(x), 2))
        else:
            try:
                yoy.append(int(x))
            except:
                error_show('Введены буквы вместо цифр в таблицах')
    return yoy
# -----------------------------------------------------------------------------
'''Функция сбора раскрывающихся списков (боксов) из списков соостветствующие ячейки таблицы'''
def tablist(NameComboBox, tablica, strok, stolbec, listZna4):
    eval('ui.{}.setObjectName("{}")'.format(NameComboBox, NameComboBox))
    for x in range(0, len(listZna4)):
        eval('ui.{}.addItem("")'.format(NameComboBox))
        eval('ui.{}.setItemText({}, _translate("Form", "{}"))'.format(NameComboBox, x, listZna4[x]))
    eval('ui.{}.setFrame(False)'.format(NameComboBox))
    eval('ui.{}.setStyleSheet("background-color: rgb(254, 254, 254);")'.format(NameComboBox))
    eval('ui.tableWidget{}.setCellWidget({}, {}, ui.{})'.format(tablica, strok, stolbec, NameComboBox))

'''НДС'''
# -----------------------------------------------------------------------------

NDSALL = [
        '       1',
        '     2, 3  ']

ui.comboBox_0 = QtWidgets.QComboBox(ui.tab_3)
tablist(NameComboBox = 'comboBox_0', tablica = '_14', strok = 3, stolbec = 0, listZna4 = NDSALL)

'''Сечение балки профиль'''
ProfALL = [
        ' Двутавр ГОСТ 8239-89 ',
        ' Швеллер У ГОСТ 8240-97',
        ' Швеллер П ГОСТ 8240-97                        '
        ]
        
ui.comboBox_1 = QtWidgets.QComboBox(ui.tab_3)
tablist(NameComboBox = 'comboBox_1', tablica = '_18', strok = 0, stolbec = 0, listZna4 = ProfALL)

'''Сечение балки номер'''
ui.comboBox_2 = QtWidgets.QComboBox(ui.tab_3)

ProDT = [
    '      10', 
    '      12', 
    '      14', 
    '      16', 
    '      18', 
    '      20', 
    '      22', 
    '      24', 
    '      27', 
    '      30', 
    '      33', 
    '      36', 
    '      40', 
    '      45', 
    '      50', 
    '      55', 
    '      60   ']

ProSHY = [
    '     5У',
    '     6.5У',
    '     8У',
    '     10У',
    '     12У',
    '     14У',
    '     16У',
    '     16аУ',
    '     18У',
    '     18аУ',
    '     20У',
    '     22У',
    '     24У',
    '     27У',
    '     30У',
    '     33У',
    '     36У',
    '     40У']

ProSHP = [
    '     5П',
    '     6.5П',
    '     8П',
    '     10П',
    '     12П',
    '     14П',
    '     16П',
    '     16аП',
    '     18П',
    '     18аП',
    '     20П',
    '     22П',
    '     24П',
    '     27П',
    '     30П',
    '     33П',
    '     36П',
    '     40П']

'''Список номеров по умолчанию для двутовров'''
tablist(NameComboBox = 'comboBox_2', tablica = '_21', strok = 0, stolbec = 0, listZna4 = ProDT)

'''Функция подстановки списка номеров по выбранному профилю'''
def DYorSH():
    if ui.comboBox_1.currentText() == ProfALL[0] and ui.toolBox.currentIndex() == 0:
        # tablist(NameComboBox = 'comboBox_2', tablica = '_21', strok = 0, stolbec = 0, listZna4 = ProDT)
        tablist(NameComboBox = 'comboBox_2', tablica = '_21', strok = 0, stolbec = 0, listZna4 = ProDT)
    if ui.comboBox_1.currentText() == ProfALL[1] and ui.toolBox.currentIndex() == 0:
        tablist(NameComboBox = 'comboBox_2', tablica = '_21', strok = 0, stolbec = 0, listZna4 = ProSHY)
    if ui.comboBox_1.currentText() == ProfALL[2] and ui.toolBox.currentIndex() == 0:
        tablist(NameComboBox = 'comboBox_2', tablica = '_21', strok = 0, stolbec = 0, listZna4 = ProSHP)

# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------
'''Очищаем раскрывающийся список номеров сечения, чтобы не оставались предыдущие выборы списка'''
ui.comboBox_1.activated['QString'].connect(ui.comboBox_2.clear)
'''По выббранному профилю сечения меняем и список номеров сечения'''
ui.comboBox_1.activated['QString'].connect(DYorSH)
# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------
'''Выбираем марку стали'''
MarkaStaList = [
    '    С345',
    '    С255',
    '    С245',
    '  другая']

ui.comboBox_3 = QtWidgets.QComboBox(ui.tab_3)
tablist(NameComboBox = 'comboBox_3', tablica = '_14', strok = 4, stolbec = 0, listZna4 = MarkaStaList)

# -----------------------------------------------------------------------------
'''Выбираем толщину стали по марке'''
ui.comboBox_3_1 = QtWidgets.QComboBox(ui.tab_3)
TolStal345i255 = [
    '    4 - 10',
    '   10 - 20',
    '   20 - 40']
TolStal245 = [
    '    4 - 20',
    '   20 - 40']
TolStallubay = [
    '   любая']

'''Список толщин стали по умолчанию'''
tablist(NameComboBox = 'comboBox_3_1', tablica = '_14', strok = 5, stolbec = 0, listZna4 = TolStal345i255)

def Still():
    '''Механические свойства стали по СП 16.1333.2017 
    Таблица В.5 - Нормативные и расчетные сопротивления при растяжении, 
    сжатии и изгибе фасонного проката
    по ГОСТ 27772 из НормКада'''
    С345_04_10 = (345, 480, 340, 470, 197, 'C345')
    С345_10_20 = (325, 470, 320, 460, 185, 'C345')
    С345_20_40 = (305, 460, 300, 450, 174, 'C345')
    С255_04_10 = (255, 380, 250, 370, 145, 'C255')
    С255_10_20 = (245, 370, 240, 360, 139, 'C255')
    С255_20_40 = (235, 370, 230, 360, 133, 'C255')
    С245_04_20 = (245, 370, 240, 360, 139, 'C245')
    С245_20_40 = (235, 370, 230, 360, 133, 'C245')
    C345ALL = [С345_04_10, С345_10_20, С345_20_40]
    C255ALL = [С255_04_10, С255_10_20, С255_20_40]
    C245ALL = [С245_04_20, С245_20_40]
    # CDrugay = []
    '''Находим расчетные значения стали'''
    # Если C345
    if ui.comboBox_3.currentText() == MarkaStaList[0]:       
        Steel = C345ALL[ui.comboBox_3_1.currentIndex()]
    # Если C255
    if ui.comboBox_3.currentText() == MarkaStaList[1]:       
        Steel = C255ALL[ui.comboBox_3_1.currentIndex()]
    # Если C245
    if ui.comboBox_3.currentText() == MarkaStaList[2]:       
        Steel = C245ALL[ui.comboBox_3_1.currentIndex()]
    # Если C - другая
    global Ry, Rs, MarkaStali
    if ui.comboBox_3.currentText() == MarkaStaList[3]:       
        try:
            Ry = ui.tableWidget_14.item(6, 0).text()
            Rs = ui.tableWidget_14.item(7, 0).text()
            if Ry == '' or Rs == '':
                error_show('Сопротивление стали не может равняться нулю')
                Ry = 0
                Rs = 0
            if Ry == 0 or Rs == 0:
                error_show('Сопротивление стали не может равняться нулю')
            Ry = float(Ry)
            Rs = float(Rs)
        except:
            error_show('Введите сопротивления стали')
    else:
        SoprName = ('Ryn', 'Run', 'Ry', 'Ru', 'Rs', 'MarkaStali')
        SteelX = {SoprName[i] : Steel[i] for i in range(0, len(SoprName))}
        # Ym = 1.025
        # # SteelX['Ry'] = round(round(SteelX['Ryn'] / Ym / 5) * 5)
        # # SteelX['Rs'] = round(round(SteelX['Ryn'] * 0.58 / Ym / 5) * 5)
        # SteelX['Ry'] = round(SteelX['Ryn'] / Ym,5)
        # SteelX['Rs'] = round(0.58 * SteelX['Ryn'] / Ym)
        Ry = SteelX['Ry']
        Rs = SteelX['Rs']
        MarkaStali = SteelX['MarkaStali']
    return Ry, Rs


'''Функция подстановки списка толщин стали по выбранной марки стали'''
_translate = QtCore.QCoreApplication.translate
def TolStlist():
    if ui.comboBox_3.currentText() == MarkaStaList[0] or ui.comboBox_3.currentText() == MarkaStaList[1]:
        tablist(NameComboBox = 'comboBox_3_1', tablica = '_14', strok = 5, stolbec = 0, listZna4 = TolStal345i255)
    if ui.comboBox_3.currentText() == MarkaStaList[2]:
        tablist(NameComboBox = 'comboBox_3_1', tablica = '_14', strok = 5, stolbec = 0, listZna4 = TolStal245)
    if ui.comboBox_3.currentText() == MarkaStaList[3]:
        tablist(NameComboBox = 'comboBox_3_1', tablica = '_14', strok = 5, stolbec = 0, listZna4 = TolStallubay)
    global Ry, Rs
    Ry = Still()[0]
    Rs = Still()[1]
    
    

    ui.tableWidget_14.item(6, 0).setText(_translate("Form", str(Ry)))
    ui.tableWidget_14.item(7, 0).setText(_translate("Form", str(Rs)))
TolStlist()
# -----------------------------------------------------------------------------
'''Очищаем раскрывающийся список номеров сечения, чтобы не оставались предыдущие выборы списка'''
ui.comboBox_3.activated['QString'].connect(ui.comboBox_3_1.clear)
'''По выббранному профилю сечения меняем и список номеров сечения'''
ui.comboBox_3.activated['QString'].connect(TolStlist)
ui.comboBox_3_1.activated['QString'].connect(TolStlist)
# -----------------------------------------------------------------------------

'''Число закреплений сжатого пояса в пролете:'''
ChisloZakrep = [
        ' без закреплений',
        ' два и более, делящие пролет балки на равные части',
        ' одно в середине       ']

'''Число загружений для составного двутавра с одной осью симметрии'''
ChisloZakrepSD = [
    ' любое {}'.format('   ' * 25)]

ui.comboBox_4 = QtWidgets.QComboBox(ui.tab_3)

'''Вид нагрузки в пролете:'''
VidNagrBezZakrep = [
        ' равномерно распределенная',
        ' сосредоточенная                                                          ']
VidNagr2ibolee =   [
        ' любая                                         ']
VidNagr1Vcentre =  [
        ' cосредоточенная в середине',
        ' сосредоточенная в четверти',
        ' равномерно распределенная                                         ']

'''Список вида нагрузки для составного двутавра с одной осью симметрии'''
VidNagrSD =  [
    ' cосредоточенная в середине',
    ' равномерно распределенная',
    ' вызывающей чистый изгиб                                          ']
ui.comboBox_5 = QtWidgets.QComboBox(ui.tab_3)

'''Функция подстановки списка вида нагрузки по выбранному числу закреплений'''
def VidNaga():
    if ui.comboBox_4.currentText() == ChisloZakrep[0]:
        tablist(NameComboBox = 'comboBox_5', tablica = '_19', strok = 3, stolbec = 0, listZna4 = VidNagrBezZakrep)
    if ui.comboBox_4.currentText() == ChisloZakrep[1]:
        tablist(NameComboBox = 'comboBox_5', tablica = '_19', strok = 3, stolbec = 0, listZna4 = VidNagr2ibolee)
    if ui.comboBox_4.currentText() == ChisloZakrep[2]:
        tablist(NameComboBox = 'comboBox_5', tablica = '_19', strok = 3, stolbec = 0, listZna4 = VidNagr1Vcentre)

def ChisloZak():
    SosDv = vvod('_33', 6, 0)
    '''Ширина верхней полки'''
    bvp = SosDv[2]
    '''Толщтна верхней полки'''
    tvp = SosDv[3]
    '''Ширина нижней полки'''
    bnp = SosDv[4]
    '''Толщтна нижней полки'''
    tnp = SosDv[5]
    if ui.toolBox.currentIndex() != 4:
        tablist(NameComboBox = 'comboBox_4', tablica = '_19', strok = 1, stolbec = 0, listZna4 = ChisloZakrep)
        VidNaga()
    if ui.toolBox.currentIndex() == 4:
        if bvp == bnp or tvp == tnp:
            tablist(NameComboBox = 'comboBox_4', tablica = '_19', strok = 1, stolbec = 0, listZna4 = ChisloZakrep)
            VidNaga()
    if ui.toolBox.currentIndex() == 4:
        if bvp != bnp or tvp != tnp:
            tablist(NameComboBox = 'comboBox_4', tablica = '_19', strok = 1, stolbec = 0, listZna4 = ChisloZakrepSD)
            tablist(NameComboBox = 'comboBox_5', tablica = '_19', strok = 3, stolbec = 0, listZna4 = VidNagrSD)
'''Выствляем значения начальные значения количества загружений'''
ChisloZak()
# -----------------------------------------------------------------------------
ui.tabWidget_4.tabBarClicked[int].connect(ui.comboBox_4.clear)
ui.tabWidget_4.tabBarClicked[int].connect(ui.comboBox_5.clear)
ui.tabWidget_4.tabBarClicked[int].connect(ChisloZak)

# ui.toolBox.currentChanged[int].connect(ui.comboBox_4.clear)
# ui.toolBox.currentChanged[int].connect(ui.comboBox_5.clear)
# ui.toolBox.currentChanged[int].connect(ChisloZak)
# -----------------------------------------------------------------------------
'''Очищаем раскрывающийся список перед вставкой вида нагрузок, чтобы не оставались предыдущие выборы списка'''
ui.comboBox_4.activated['QString'].connect(ui.comboBox_5.clear)
'''По выббранному числу закреплений подставить вид нагрузки'''
ui.comboBox_4.activated['QString'].connect(VidNaga)
# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------

'''Список квадратных сечений размещаем в бокс'''
KvadList = Sortament.ListElement('kvadrat')
KvadListName = [KvadList[i]['Name'] for i in range(0, len(KvadList))]
ui.comboBox_6 = QtWidgets.QComboBox(ui.tab_3)
tablist(NameComboBox = 'comboBox_6', tablica = '_31', strok = 0, stolbec = 0, listZna4 = KvadListName)

'''Список прямоугольных сечений размещаем в бокс'''
PraymoygList = Sortament.ListElement('praymoygol')
PraymoygListName = [PraymoygList[i]['Name'] for i in range(0, len(PraymoygList))]
ui.comboBox_7 = QtWidgets.QComboBox(ui.tab_3)
tablist(NameComboBox = 'comboBox_7', tablica = '_37', strok = 0, stolbec = 0, listZna4 = PraymoygListName)


# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------

# # Проверка ячеек на пустые значения и ","
# def vvod(nomerwidgeta, strok, stolbec):
#     yoy = []
#     for i in range(1, strok+1):
#         x = eval('ui.tableWidget{}.item({}-1, {}).text()'.format(nomerwidgeta, i, stolbec))
#         if x != '': 
#             x = x.replace(',', '.')
#         else: 
#             x = 0
#         if '.' in str(x):
#             yoy.append(round(float(x), 2))
#         else:
#             try:
#                 yoy.append(int(x))
#             except:
#                 error_show('Введены буквы вместо цифр в таблицах')
#     return yoy

def delet_0(x):
    try:
        while x[-1] == 0: del x[-1]
    except:
        x.append(0)

def sbor_dannih():
    global Lkl, Lbk, Lkp, Nv, Nvx1, Nvx2, Moment, MomX, dannie, box_index
    # sechenie = ui.radioButton.isChecked()

    Lkl = vvod('_20', 1, 0)          # длина левой консоли балки
    Lbk = vvod('_20', 1, 1)          # длина балки без консоли
    Lkp = vvod('_20', 1, 2)          # длина правой консоли балки

    Nv = vvod('_11', 20, 0)          # значения вертикальных нагрузок
    Nvx1 = vvod('_11', 20, 1)        # начало вертикальных нагрузок
    Nvx2 = vvod('_11', 20, 2)        # конец вертикальных нагрузок

    Moment = vvod('_12', 20, 0)      # значения моментов
    MomX = vvod('_12', 20, 1)        # привязка моментов

    dannie = vvod('_14', 8, 0)       # данные с таблицы

    """Собираем индексы с боксов"""
    '''
    ui.comboBox_0 - НДС
    ui.comboBox_1 - Сечение балки профиль
    ui.comboBox_2 - Сечение балки номер
    ui.comboBox_3 - Марка стали
    ui.comboBox_4 - Число закреплений сжатого пояса в пролете
    ui.comboBox_5 - Вид нагрузки в пролете
    ui.comboBox_3_1 - Толщина стали
    '''
    box_index = []
    for i in range(0, 6):
        y = eval('ui.comboBox_{}.currentIndex()'.format(i))
        box_index.append(y)
    x = ui.comboBox_3_1.currentIndex()
    box_index.append(x)
    
# block()
# -----------------------------------------------------------------------------
def raschet():
    
    _translate = QtCore.QCoreApplication.translate
    ui.label_3.setText(_translate("Form", "Отчет"))
    global Lkl, Lbk, Lkp, Nv, Nvx1, Nvx2, Moment, MomX, dannie, Ry, Rs, MarkaStali
    global document
    global G, Ystouchnvost, RfR, fu, fRR, ful, RRf, fup, Tay, alfa, Stop841
    # global noyblock
    # global RfR, fu
    del document    # удаляем глобальную переменную и заменяем ее новой такой же, но пустой
    document = Document()
    
    TolStlist()

    if Ry == 0 or Rs == 0:
        error_show('Сопротивление стали не может равняться нулю')
        return
    
    sbor_dannih()

    delet_0(Nv)
    delet_0(Moment)
    
    # 

    # Корректировка спиков по длине спика Nv
    kk = len (Nv); Nvx1 = Nvx1 [:kk]; Nvx2 = Nvx2 [:kk]
    # Корректировка спиков по длине спика Moment
    kk = len (Moment); MomX = MomX [:kk]
    # Длина балки
    L = sum(Lkl + Lbk + Lkp)
    # Координаты опорных реакций
    xR1 = 0 if Lkl == 0 else sum(Lkl)
    xR2 = sum(Lkl + Lbk)
    
    ''' Подготовка списков для расчета опорных реакций R1 и R2: '''
    NvP = []
    Nvx = []

    for i in range(0, len(Nvx2)):
        if Nvx2[i] != 0:
            Nvx.append(((Nvx2[i] - Nvx1[i])/2) + Nvx1[i] - xR1)
            NvP.append(Nv[i] * (Nvx2[i] - Nvx1[i]))
        else:
            Nvx.append(Nvx1[i] - xR1)
            NvP.append(Nv[i])

    M = Moment
    xM = MomX

    '''
    Правило задание моментов при нахождении R2
    момент (+) - вращении по часовой стрелки.
    момент (-) - вращении против часовой стрелки.
    '''
    # Сумма моментов в точке приложения R1:
    R2 = []
    if xR1 == 0:
        for i in range(0, len(NvP)):
            try:
                R2.append(NvP[i] * Nvx[i] / xR2)
            except:
                error_show('Укажите длину балки')
                return
        for i in range(0, len(M)):
            R2.append(-M[i] / xR2)
    
    if xR1 > 0:
        for i in range(0, len(NvP)):
            R2.append(NvP[i] * Nvx[i] / (xR2 - xR1))
        for i in range(0, len(M)):
            R2.append(-M[i] / (xR2 - xR1))

    R2 = -round(sum(R2), 5)
    R1 = -round(sum(NvP) + R2, 5)
    
    

    ''' Привязки всех нагрузок'''
    privyazki = Nvx1 + Nvx2 + xM + [xR1] + [xR2] + [L]
    # 
    privyazki = sorted(privyazki)
    listpriv = [0]
    for i in range(1, len(privyazki)):
        if privyazki[i] != privyazki[i-1]:
            listpriv.append(privyazki[i])
        else:
            continue
    privyazki = listpriv
    Ychastki = []
    for i in range(1, len(privyazki)):
        Ychastki.append([privyazki[i-1], privyazki[i]])

    '''Создаем списки координат и нагрузок для нахождения моментов и поперечной силы'''
    xP = []             # координаты сосредоточенных нагрузок 
    xQ = []             # координаты распределенных нагрузок
    xQ0 = []            # координаты распределенных нагрузок для отображения схемы кратные 0,1м
    LQ = []             # протяженность распределенных нагрузок
    P = []              # значения сосредоточенных нагрузок
    Q = []              # значения распределенных нагрузок

    # 

    for i in range(0, len(Nvx2)):
        if Nvx2[i] == 0:
            P.append(-Nv[i])
            xP.append(Nvx1[i])
        else:
            Q.append(Nv[i])
            x = Nvx1[i]
            y = []
            y.append(x)
            while x < Nvx2[i]:
                x = round(x + 0.01, 2)
                y.append(x)
            xQ.append(y)
            z = 0.0
            yl = []
            yl.append(z)
            while z < round(Nvx2[i] - Nvx1[i], 2):
                z = round(z + 0.01, 2)
                yl.append(z)
            LQ.append(yl)

    for i in range(0, len(Nvx2)):
        if Nvx2[i] != 0:
            x = Nvx1[i]
            y = []
            y.append(x)
            while x < Nvx2[i]:
                x = round(x + 0.1, 3)
                y.append(x)
            xQ0.append(y)

    for i in range(0, len(xQ0)):
        del xQ0[i][-1]
        xQ0[i].append(xQ[i][-1])

    xP0 = xP[:]
    P0 = P[:]

    '''Дополняем сосредоточенные нагрузки опорной реакцией R1, R2'''
    P.append(-R1)
    P.append(-R2)
    xP.append(xR1)
    xP.append(xR2)
    
    '''Разбиваем балку по учассткам с шагом "n" '''
    # x = 0
    XL = []
    for i in range(0, len(Ychastki)):
        sss = []
        x = Ychastki[i][0]
        sss.append(x)
        while x < Ychastki[i][-1]:
            x = round(x + 0.01, 2)
            sss.append(x)
        XL.append(sss)
    # XL_ych = XL[:]
    # 

    

    # 
    # 
    # 
    # 
    # 
    # 
    # 
    # 
    # 

    '''//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////'''
    '''
    Расчет прогибов ведем по Правилу Клебша:
    1) Мх выражаем через внешние силы, которые лежат только слева или тольло справа.
    2) Если погонная сила q не доходит до правого конца, то ее доводим до этого проаваго конца и уравновешиваем ее снизу.
    3) Если имеется сосредоточенный момент Mo, то его вклад в изгибающий момент записываем в виде Mo(x - a)*, где а - расстояние до момента, * - нулевая степень.
    4) Интегрируем не раскрывая скобок.
    '''
    '''Создаем новые списки координат для каждой распределенной нагрузки от начальной точки нагрузки до конца балки'''
    xQL = []
    LQL = []
    for i in range(0, len(Nvx2)):
        if Nvx2[i] != 0:
            x = Nvx1[i]
            y = []
            y.append(x)
            while x < L:
                x = round(x + 0.01, 2)
                y.append(x)
            xQL.append(y)
            z = 0.0
            yl = []
            yl.append(z)
            while z < round(L - Nvx1[i], 2):
                z = round(z + 0.01, 2)
                yl.append(z)
            LQL.append(yl)
    # 
    # 

    '''Создаем новые списки координат для каждой распределенной нагрузки от конечной точки нагрузки до конца балки'''
    x2QL = []
    L2QL = []
    for i in range(0, len(Nvx2)):
        if Nvx2[i] != 0:
            x = Nvx2[i]
            y = []
            y.append(x)
            while x < L:
                x = round(x + 0.01, 2)
                y.append(x)
            x2QL.append(y)
            z = 0.0
            yl = []
            yl.append(z)
            while z < round(L - Nvx2[i], 2):
                z = round(z + 0.01, 2)
                yl.append(z)
            L2QL.append(yl)
    # 
    # 

    
    '''Находим известные данные от двойного интегрирования моментов в точках с нулевым прогибом xR1 и xR2'''
    '''EJvII = - M;   EJ0 = EJvI'''

    # 
    def doubleintegral():
        yyy = []
        for i in range(0, len(XL)):
            xxxx = []
            for x in range(0, len(XL[i])):
                '''# перебираем координаты сосредоточенных нагрузок и складываем их значения согласно X:'''
                vP = 0
                for q in range(0, len(P)) or XL[i][0] == xP[q]:
                    if XL[i][x] > xP[q]:
                        vP = vP + (P[q] * (XL[i][x] - xP[q]) ** 3) / 6
                '''# перебираем координаты распределенных нагрузок и складываем их значения согласно X:'''
                vQ = 0
                vQl = 0
                for w in range(0, len(Q)):
                    for e in range(0, len(xQL[w])):
                        if XL[i][x] == xQL[w][e]:
                            # перебираем координаты для каждой распределенной нагрузки от НАЧАЛЬНОЙ точки нагрузки до конца балки
                            vQ = vQ + (Q[w] * LQL[w][e] ** 4) / 24
                    for k in range(0, len(L2QL[w])):
                        if XL[i][x] == x2QL[w][k]:
                            # перебираем координаты для каждой распределенной нагрузки от КОНЕЧНОЙ точки нагрузки до конца балки со знаком (-)
                            vQl = vQl + (Q[w] * L2QL[w][k] ** 4) / 24
                '''# перебираем координаты нагрузок от приложенных моментов'''
                vj = 0
                for g in range(0, len(M)):
                    if XL[i][x] > xM[g] or XL[i][0] == xM[g]:
                        vj = vj + (M[g] * XL[i][x] ** 2 * 0.5)
                try:
                    xxxx.append(round( - vP + vQ - vQl + vj + C * XL[i][x] + D, 5))
                except:
                    xxxx.append(round( - vP + vQ - vQl + vj, 5))
            yyy.append(xxxx)
        return yyy

    '''Находим постоянные С и D на опорах с нулевым прогибом'''
    EJvIIALL = doubleintegral()
    
    EJvxR1xR2 = []
    for i in range(0, len(XL)):
        for x in range(0, len(XL[i])):
            if XL[i][x] == xR1 or XL[i][x] == xR2:
                s = EJvIIALL[i][x]
                EJvxR1xR2.append(s)
    EJvxR1xR2 = [EJvxR1xR2[0], EJvxR1xR2[-1]] 

    C = - (EJvxR1xR2[1] - EJvxR1xR2[0])/(xR2 - xR1)
    D = - C * xR1 - EJvxR1xR2[0]
# ==========================================================================================================

    '''Формируем списки от двойного интегрирования моментов с учетом С и D для эпюры прогибов'''
    EJvII = doubleintegral()

    '''Находим максимальные значения прогибов по участкам балки'''
    EJvII_I = []
    XLEJvII = []
    for i in range(0, len(EJvII)):
        if max(EJvII[i]) != 0:
            EJvII_I.append(round(max(EJvII[i]), 3))
            XLEJvII.append(XL[i][EJvII[i].index(max(EJvII[i]))])
        if min(EJvII[i]) != 0:
            EJvII_I.append(round(min(EJvII[i]), 3))
            XLEJvII.append(XL[i][EJvII[i].index(min(EJvII[i]))])

    '''Собираем все участки в один список для построения эпюры EJvf'''
    EJvII0 = []
    for i in range(0, len(EJvII)):
        for x in range(0, len(EJvII[i])):
            EJvII0.append(EJvII[i][x])
    EJvII = EJvII0
    EJvIImin = min(EJvII)
    EJvIImax = max(EJvII)
    # EJvIImin = round(min(EJvII),5)
    # EJvIImax = round(max(EJvII),5)
    # EJvIIabsmax = max(abs(EJvIImax), abs(EJvIImin))
    
 



# ==========================================================================================================
    '''//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////'''

    '''Находим значения эпюры Mx и Qy в сечениях XL'''
    # Для каждого сечения XL
    if xP[-1] == L:
        del xP[-1]
        del P[-1]
    M_x = []; Q_x = []
    for i in range(0, len(XL)):
        mmmm = []; qqqq = []
        for x in range(0, len(XL[i])):
            y = 0; a = 0
            '''# перебираем координаты сосредоточенных нагрузок'''
            for q in range(0, len(P)):
                # берем только значению меньше XL не вулючая нагрузку в конце участка и беря ее в начале следующего участка
                if XL[i][x] > xP[q] or XL[i][0] == xP[q]:
                    # умножаем значение силы на плечо XL, складиываю при этом с предыдущим вычислением
                    y = y + P[q] * (XL[i][x] - xP[q])
                    a = round(a + P[q], 5)
            '''# перебираем координаты для каждой распределенной нагрузки'''
            u = 0; b = 0
            for w in range(0, len(Q)):
                for e in range(0, len(xQ[w])):
                    # если распределенная нагрузка доходит до конца балки
                    if XL[i][x] == xQ[w][e]:
                        # умножаем значение силы на плечо XL, складиываю при этом с предыдущим вычислением
                        u = u + Q[w] * LQ[w][e] * (LQ[w][e] * 0.5)
                        b = b + Q[w] * LQ[w][e]
                # если распределенная нагрузка НЕ доходит до конца балки
                if XL[i][x] > xQ[w][-1]:
                    # умножаем значение силы на плечо XL, складиываю при этом с предыдущим вычислением
                    u = round(u + Q[w] * LQ[w][-1] * ((XL[i][x] - xQ[w][-1]) + LQ[w][-1] * 0.5), 5)
                    b = round(b + Q[w] * LQ[w][-1], 5)
            '''# перебираем координаты нагрузок от приложенных моментов'''
            j = 0
            for g in range(0, len(M)):
                if XL[i][x] > xM[g] or XL[i][0] == xM[g]:
                    j = j + M[g]
            '''Формируем слагаемые в значения по X'''
            # M_x.append(round(y - u - j, 5))
            # Q_x.append(round(a - b, 5))
            
            mmmm.append(round(y - u - j, 5))
            qqqq.append(round(a - b, 5))
        M_x.append(mmmm)
        Q_x.append(qqqq)

        '''Конец перебора'''

    '''Собираем значения с эпюр Мх и Qу на границах участков'''
    granQI = []; granMI = []; granXI = []
    # for i in range(0, len(XL)):
    for i in range(0, len(XL)):
        x = []; y = []; z = []
        x.append(round(Q_x[i][0], 3))
        x.append(round(Q_x[i][-1], 3))
        granQI.append(x)
        y.append(round(M_x[i][0], 3))
        y.append(round(M_x[i][-1], 3))
        granMI.append(y)
        z.append(round(XL[i][0], 3))
        z.append(round(XL[i][-1], 3))
        granXI.append(z)

    
    # 
    


    '''Находим максимальные значения  Mx и Qy по участкам балки'''
    Mx_max_min = []
    XL_Mx_max_min = []
    # Q_x_max_min = []
    # XL_Qx_max_min = []
    for i in range(0, len(XL)):
        if max(M_x[i]) != 0 and max(M_x[i]) not in Mx_max_min:
            Mx_max_min.append(round(max(M_x[i]), 3))
            XL_Mx_max_min.append(XL[i][M_x[i].index(max(M_x[i]))])
        if min(M_x[i]) != 0 and min(M_x[i]) not in Mx_max_min:
            Mx_max_min.append(round(min(M_x[i]), 3))
            XL_Mx_max_min.append(XL[i][M_x[i].index(min(M_x[i]))])
        # if max(Q_x[i]) != 0 and max(Q_x[i]) not in Q_x_max_min:
        #     Q_x_max_min.append(round(max(Q_x[i]), 3))
        #     XL_Qx_max_min.append(XL[i][Q_x[i].index(max(Q_x[i]))])
        # if min(Q_x[i]) != 0 and min(Q_x[i]) not in Q_x_max_min:
        #     Q_x_max_min.append(round(min(Q_x[i]), 3))
        #     XL_Qx_max_min.append(XL[i][Q_x[i].index(min(Q_x[i]))])
    # 
    # 
    # 
    # 





    
    # 
    '''Собираем все участки в один список для построения эпюры Mx и Qy'''
    M_x_ychastoki = []; Q_x_ychastoki = []
    for i in range(0, len(XL)):
        for x in range(0, len(XL[i])):
            M_x_ychastoki.append(M_x[i][x])
            Q_x_ychastoki.append(Q_x[i][x])
    M_x = M_x_ychastoki
    Q_x = Q_x_ychastoki

    Mmax = max(M_x)
    Mmin = min(M_x)
    Qmax = max(Q_x)
    Qmin = min(Q_x)

    '''Координаты сечений X перегоняем в один список для построения эпюр'''
    XLL = []
    for i in range(0, len(XL)):
        for x in range(0, len(XL[i])):
            XLL.append(XL[i][x])
    XL = XLL

    '''Собираем значения с эпюр Мх и Qу на границах участков в один список'''
    granQ = []; granM = []; granX = []
    for i in range(0, len(XL)):
        for e in range(0, len(privyazki)):
            if XL[i] == privyazki[e]:
                granQ.append(round(Q_x[i], 3))
                granM.append(round(M_x[i], 3))
                granX.append(round(XL[i], 3))
    granX0 = granX[:]
    # granM0 = granM[:]
    # granQ0 = granQ[:]

    '''Создаем список со значениями максимум и минимиум с каждого участка с вычетом значений на границах'''
    MImax = []
    xImax = []
    mm = []; xx = []
    for i in range(0, len(Mx_max_min)):
        if Mx_max_min[i] not in granM:
            # granQ.append(0.0)
            mm.append(round(Mx_max_min[i], 3))
            xx.append(round(XL_Mx_max_min[i], 3))
    MImax.append(mm)
    xImax.append(xx)

    # 
    # 

    '''Добавляем в список со значениями границ участков максимумы и минимиумы с каждого участка'''
    for i in range(0, len(Mx_max_min)):
        if Mx_max_min[i] not in granM:
            granQ.append(0.0)
            granM.append(round(Mx_max_min[i], 3))
            granX.append(round(XL_Mx_max_min[i], 3))

    # 
    # 
    # 

    '''Список максимльных прогибов между опорами'''
    EJvIIR = EJvII[XL.index(xR1):XL.index(xR2)+1]
    EJvIIRmax = round(max(EJvIIR), 5)
    EJvIIRmin = round(min(EJvIIR), 5)
    '''Максимальное значение прогиба между опорами по модулю'''
    EJvIIRabs = round( max(abs(EJvIIRmax), abs(EJvIIRmin)) , 5)
    '''Максимальное значение прогиба консоли по модулю'''
    if sum(Lkl) != 0.0:
        EJvIIK1max = round(max(EJvII[ : XL.index(xR1)]), 5)
        EJvIIK1min = round(min(EJvII[ : XL.index(xR1)]), 5)
        EJvIIK1abs = max(abs(EJvIIK1max), abs(EJvIIK1min))
        # EJvIIKmax = EJvIIK1abs
    if sum(Lkp) != 0.0:
        EJvIIK2max = round(max(EJvII[XL.index(xR2)+1 : ]), 5)
        EJvIIK2min = round(min(EJvII[XL.index(xR2)+1 : ]), 5)
        EJvIIK2abs = max(abs(EJvIIK2max), abs(EJvIIK2min))
        # EJvIIKmax = EJvIIK2abs
    # if Lkl != 0 and Lkp != 0:
    #     EJvIIKmax = max(EJvIIK1abs, EJvIIK2abs)

    # 
    # 
    # 
    # 
    


    '''====================================================================================================================================='''
    '''====================================================================================================================================='''
    class Example(QWidget):
        def __init__(self):
            super().__init__()
            self.initUI()
        def initUI(self):
            self.setGeometry(0, 0, Wrisunka, Hrisunka)
            self.setStyleSheet("background-color: rgb(0, 0, 0, 0);")
        def paintEvent(self, e):
            qp = QPainter()
            qp.begin(self)
            self.drawBrushes(qp)
            qp.end()
        def drawBrushes(self, qp):
            '''Граничные точки по углам и снизу в центре для ориентира'''
            # qp.setPen(QPen(QtGui.QColor(0, 100, 150), 4, QtCore.Qt.SolidLine))
            # qp.drawPoint(0, 0)
            # qp.drawPoint(0, Hrisunka)
            # qp.drawPoint(Wrisunka, 0)
            # qp.drawPoint(Wrisunka/2, Hrisunka)
            # qp.drawPoint(Wrisunka, Hrisunka)
            
            '''qp.draWТext(<X>, <У>, <Ширина>, <Высота>, <Флаги>, <Текст>)'''
            '''AlignLeft, AlignRight, AlignHCenter, AlignVCenter, AlignCenter, (TextDontClip - не резать рамками текст)'''
            # brush = QBrush(Qt.NoBrush)
            # brush.setColor(QtGui.QColor(0,100,150,150))
            # brush.setStyle(Qt.CrossPattern)
            # qp.setBrush(brush)
            # qp.drawRect(0, 0, 50, 30)
            # qp.drawText(0, 0, 50, 30, Qt.AlignLeft, '4')
            # -------------------------------------------------------------------------
            '''Сдвигаем координаты Х'''
            dx = 15

            ''' масштаб эпюр по горизонту '''
            # masshtabX = (ui.frame_12.width() - 40)/L
            masshtabX = (Wrisunka - 30)/L

            '''координата оси балки'''
            dy = 40
            if len(P0) == 0 or sum(M) == 0:
                nP = 0
            if len(P0) != 0 or sum(M) != 0:
                nP = 1
            nN = len(Q) + nP
            osBalki = dy * nN + 20

            '''Максимальное количество пикселей по вертикали на каждую эпюру
            с вычетом места под надписи эпюр и цифр значений:
            90 + 45 + 60 + 60 + 15 = 270
            90 - расстояние от оси балки до нижней точки размерной линии
            60 + 60 - расстояния между крайними точками соседних эпюр
            15 - расстояние от нижней точки последней эпюры до конца рисунка

            90 + 50 + 70 + 70 + 15 = 295
            90 - расстояние от оси балки до нижней точки размерной линии
            60 + 60 - расстояния между крайними точками соседних эпюр
            15 - расстояние от нижней точки последней эпюры до конца рисунка'''
            # Yos = (Hrisunka - osBalki - 270)/3 - 0
            Yos = (Hrisunka - osBalki - 295)/3 - 0

            '''Согласно выделенного количество пикселей по вертикали на каждую эпюру
            масштабируем вертикальные значения эпюр'''
            try:
                masshtabYQ = Yos/(abs(Qmax) + abs(Qmin))
            except:
                masshtabYQ = 1
            try:
                masshtabYM = Yos/(abs(Mmax) + abs(Mmin))
            except:
                masshtabYM = 1
            try:
                masshtabYEJvII = Yos/(abs(EJvIImax) + abs(EJvIImin))
            except:
                masshtabYEJvII = 1

            '''Координата оси эпюры Qy'''
            osQy = abs(Qmax) * masshtabYQ + osBalki + 140

            '''Координата оси эпюры Mx'''
            # osMx = osQy + abs(Qmin) * masshtabYQ + 60 + abs(Mmin) * masshtabYM
            osMx = osQy + abs(Qmin) * masshtabYQ + 70 + abs(Mmin) * masshtabYM

            '''Координата оси эпюры EJvII'''
            # osEJvII = osMx + Mmax * masshtabYM + 60 + abs(EJvIImin) * masshtabYEJvII
            osEJvII = osMx + Mmax * masshtabYM + 70 + abs(EJvIImin) * masshtabYEJvII



            ''' Схема загружения балки '''

            '''Опорные шарниры'''
            d = 8
            qp.setPen(QPen(QtGui.QColor(0, 100, 150), 1.5, QtCore.Qt.SolidLine))
            # drawEllipse(int x, int y, int width, int height)
            qp.drawEllipse (xR1 * masshtabX + dx - d/2, osBalki + 5, d, d)
            qp.drawEllipse (xR2 * masshtabX + dx - d/2, osBalki + 5, d, d)
            '''Наклонные линии R1 и вертикальная R2'''
            qp.drawLine(xR1 * masshtabX + dx + 3, osBalki + d + 5, xR1 * masshtabX + dx + 10, osBalki + 25)
            qp.drawLine(xR1 * masshtabX + dx - 3, osBalki + d + 5, xR1 * masshtabX + dx - 10, osBalki + 25)
            qp.drawLine(xR2 * masshtabX + dx, osBalki + d + 5, xR2 * masshtabX + dx, osBalki + 25)
            '''Горизонтальные толстые линии R1 и R2'''
            qp.setPen(QPen(QtGui.QColor(0, 100, 150), 3, QtCore.Qt.SolidLine))
            qp.drawLine(xR1 * masshtabX, osBalki + 25, xR1 * masshtabX + dx + dx, osBalki + 25)
            qp.drawLine(xR2 * masshtabX + dx - 10, osBalki + 25, xR2 * masshtabX + dx + 10, osBalki + 25)
            '''Горизонтальные тонкие линии R1 и R2'''
            qp.setPen(QPen(QtGui.QColor(0, 100, 150), 1, QtCore.Qt.SolidLine))
            qp.drawLine(xR1 * masshtabX, osBalki + 30, xR1 * masshtabX + dx + dx, osBalki + 30)
            qp.drawLine(xR2 * masshtabX + dx - 10, osBalki + 30, xR2 * masshtabX + dx + 10, osBalki + 30)

            '''Балка'''
            qp.setPen(QPen(QtGui.QColor(0, 0, 0), 3, QtCore.Qt.SolidLine))
            qp.drawLine(XL[0] * masshtabX + dx, osBalki, XL[-1] * masshtabX + dx, osBalki)
            qp.setPen(QPen(QtGui.QColor(0, 100, 150), 2, QtCore.Qt.SolidLine))
            
            '''Расстановка R1 и R2 (стрелки и значения)'''
            qp.setPen(QPen(QtGui.QColor(255, 85, 0), 3, QtCore.Qt.SolidLine))
            qp.drawLine(xR1 * masshtabX + dx, osBalki + 35, xR1 * masshtabX + dx, osBalki + 60)
            qp.setFont(QtGui.QFont("Cambria", 11))
            qp.drawText(xR1 * masshtabX + dx + 5, osBalki + 45, 50, 30, Qt.AlignLeft, str(abs(round(R1, 3))))
            qp.drawText(xR2 * masshtabX + dx - 55, osBalki + 45, 50, 30, Qt.AlignRight, str(abs(round(R2, 3))))
            qp.setPen(QPen(QtGui.QColor(255, 85, 0), 2, QtCore.Qt.SolidLine))
            if R1 <= 0:
                qp.drawLine(xR1 * masshtabX + dx, osBalki + 35, xR1 * masshtabX + dx - 3, osBalki + 43)
                qp.drawLine(xR1 * masshtabX + dx, osBalki + 35, xR1 * masshtabX + dx + 3, osBalki + 43)
                qp.drawLine(xR1 * masshtabX + dx - 3, osBalki + 43, xR1 * masshtabX + dx + 3, osBalki + 43)
            else:
                qp.drawLine(xR1 * masshtabX + dx, osBalki + 60, xR1 * masshtabX + dx - 3, osBalki + 52)
                qp.drawLine(xR1 * masshtabX + dx, osBalki + 60, xR1 * masshtabX + dx + 3, osBalki + 52)
                qp.drawLine(xR1 * masshtabX + dx - 3, osBalki + 52, xR1 * masshtabX + dx + 3, osBalki + 52)
            qp.setPen(QPen(QtGui.QColor(255, 85, 0), 3, QtCore.Qt.SolidLine))
            qp.drawLine(xR2 * masshtabX + dx, osBalki + 35, xR2 * masshtabX + dx, osBalki + 60)
            qp.setPen(QPen(QtGui.QColor(255, 85, 0), 2, QtCore.Qt.SolidLine))
            if R2 <= 0:
                qp.drawLine(xR2 * masshtabX + dx, osBalki + 35, xR2 * masshtabX + dx - 3, osBalki + 43)
                qp.drawLine(xR2 * masshtabX + dx, osBalki + 35, xR2 * masshtabX + dx + 3, osBalki + 43)
                qp.drawLine(xR2 * masshtabX + dx - 3, osBalki + 43, xR2 * masshtabX + dx + 3, osBalki + 43)
            else:
                qp.drawLine(xR2 * masshtabX + dx, osBalki + 60, xR2 * masshtabX + dx - 3, osBalki + 52)
                qp.drawLine(xR2 * masshtabX + dx, osBalki + 60, xR2 * masshtabX + dx + 3, osBalki + 52)
                qp.drawLine(xR2 * masshtabX + dx - 3, osBalki + 52, xR2 * masshtabX + dx + 3, osBalki + 52)
            
            '''Расстановка сосредоточенных нагрузок'''
            yoy = 0
            if len(P0) != 0:
                for i in range(0, len(xP0)):
                    if P0[i] != 0:
                        qp.setPen(QPen(QtGui.QColor(85, 0, 255), 3, QtCore.Qt.SolidLine))
                        qp.drawLine(xP0[i] * masshtabX + dx, osBalki - yoy - 5, xP0[i] * masshtabX + dx, osBalki - yoy - 35)
                        qp.setFont(QtGui.QFont("Cambria", 11))
                        # Зеркалим текст относительно стрелки, если он приближается на 0,5 м к краю балки, чтобы он не уходил за размеры рисунка
                        if xP0[i] < L - 0.5:
                            qp.drawText(xP0[i] * masshtabX + dx + 5, osBalki - yoy - 30, 50, 30, Qt.AlignLeft, str(abs(P0[i])))
                        else:
                            qp.drawText(xP0[i] * masshtabX + dx - 55, osBalki - yoy - 30, 50, 30, Qt.AlignRight, str(abs(P0[i])))
                        qp.setPen(QPen(QtGui.QColor(85, 0, 255), 2, QtCore.Qt.SolidLine))
                        if P0[i] <= 0:
                            qp.drawLine(xP0[i] * masshtabX + dx, osBalki - yoy - 5, xP0[i] * masshtabX + dx - 3, osBalki - yoy - 13)
                            qp.drawLine(xP0[i] * masshtabX + dx, osBalki - yoy - 5, xP0[i] * masshtabX + dx + 3, osBalki - yoy - 13)
                            qp.drawLine(xP0[i] * masshtabX + dx - 3, osBalki - yoy - 13, xP0[i] * masshtabX + dx + 3, osBalki - yoy - 13)
                        if P0[i] > 0:
                            qp.drawLine(xP0[i] * masshtabX + dx, osBalki - yoy - 35, xP0[i] * masshtabX + dx - 3, osBalki - yoy - 27)
                            qp.drawLine(xP0[i] * masshtabX + dx, osBalki - yoy - 35, xP0[i] * masshtabX + dx + 3, osBalki - yoy - 27)
                            qp.drawLine(xP0[i] * masshtabX + dx - 3, osBalki - yoy - 27, xP0[i] * masshtabX + dx + 3, osBalki - yoy - 27)

            '''Расстановка моментов'''
            if len(M) != 0:
                for i in range(0, len(M)):
                    if M[i] != 0:
                        qp.setPen(QPen(QtGui.QColor(255, 170, 0), 3, QtCore.Qt.SolidLine))
                        qp.setFont(QtGui.QFont("Cambria", 11))
                        qp.drawText(xM[i] * masshtabX + dx - 25, osBalki - yoy - 44, 50, 30, Qt.AlignHCenter, str(abs(M[i])))
                        # qp.drawText(xM[i] * masshtabX + dx + 5, osBalki - yoy - 30, 50, 30, Qt.AlignLeft, str(abs(M[i])))
                        qp.setPen(QPen(QtGui.QColor(255, 170, 0), 2, QtCore.Qt.SolidLine))
                        qp.drawLine(xM[i] * masshtabX + dx, osBalki - yoy - 20, xM[i] * masshtabX + dx, osBalki - yoy + 20)
                        if M[i] < 0:
                            qp.drawLine(xM[i] * masshtabX + dx, osBalki - yoy - 20, xM[i] * masshtabX + dx + 15, osBalki - yoy - 20)
                            qp.drawLine(xM[i] * masshtabX + dx + 7, osBalki - yoy - 23, xM[i] * masshtabX + dx + 15, osBalki - yoy - 20)
                            qp.drawLine(xM[i] * masshtabX + dx + 7, osBalki - yoy - 17, xM[i] * masshtabX + dx + 15, osBalki - yoy - 20)

                            qp.drawLine(xM[i] * masshtabX + dx, osBalki - yoy + 20, xM[i] * masshtabX + dx - 15, osBalki - yoy + 20)
                            qp.drawLine(xM[i] * masshtabX + dx - 7, osBalki - yoy + 23, xM[i] * masshtabX + dx - 15, osBalki - yoy + 20)
                            qp.drawLine(xM[i] * masshtabX + dx - 7, osBalki - yoy + 17, xM[i] * masshtabX + dx - 15, osBalki - yoy + 20)

                        if M[i] > 0:
                            qp.drawLine(xM[i] * masshtabX + dx, osBalki - yoy + 20, xM[i] * masshtabX + dx + 15, osBalki - yoy + 20)
                            qp.drawLine(xM[i] * masshtabX + dx + 7, osBalki - yoy + 23, xM[i] * masshtabX + dx + 15, osBalki - yoy + 20)
                            qp.drawLine(xM[i] * masshtabX + dx + 7, osBalki - yoy + 17, xM[i] * masshtabX + dx + 15, osBalki - yoy + 20)

                            qp.drawLine(xM[i] * masshtabX + dx, osBalki - yoy - 20, xM[i] * masshtabX + dx - 15, osBalki - yoy - 20)
                            qp.drawLine(xM[i] * masshtabX + dx - 7, osBalki - yoy - 23, xM[i] * masshtabX + dx - 15, osBalki - yoy - 20)
                            qp.drawLine(xM[i] * masshtabX + dx - 7, osBalki - yoy - 17, xM[i] * masshtabX + dx - 15, osBalki - yoy - 20)

            if sum(P0) != 0 or sum(M) != 0:
                yoy = yoy + 40

            '''Расстановка распределенных нагрузок'''
            for i in range(0, len(xQ0)):
                if Q[i] != 0:
                    qp.setPen(QPen(QtGui.QColor(0, 170, 127), 1, QtCore.Qt.SolidLine))
                    qp.setFont(QtGui.QFont("Cambria", 11))
                    qp.drawText(((xQ[i][-1] - xQ[i][0])/2 + xQ[i][0]) * masshtabX + dx - 25, osBalki - yoy - 42, 50, 20, Qt.AlignHCenter, str(Q[i]))
                    qp.drawLine(xQ0[i][0] * masshtabX + dx, osBalki - yoy - 22, xQ0[i][-1] * masshtabX + dx, osBalki - yoy - 22)
                    for x in range(0, len(xQ0[i])):
                        qp.drawLine(xQ0[i][x] * masshtabX + dx, osBalki - yoy - 5, xQ0[i][x] * masshtabX + dx, osBalki - yoy - 22)
                    yoy = yoy + 40

            '''Расстановка размеров'''
            qp.setPen(QPen(QtGui.QColor(120, 120, 120), 1, QtCore.Qt.SolidLine))
            qp.drawLine(XL[0] * masshtabX + dx - 5, osBalki + 85, XL[-1] * masshtabX + dx + 5, osBalki + 85)
            for i in range(0, len(granX0)):
                qp.drawLine(granX0[i] * masshtabX + dx - 5, osBalki + 85 + 5, granX0[i] * masshtabX + dx + 5, osBalki + 85 - 5)
                qp.drawLine(granX0[i] * masshtabX + dx, osBalki + 85 + 5, granX0[i] * masshtabX + dx, osBalki + 85 - 15)
            for i in range(0, len(Ychastki)):
                qp.drawText(((Ychastki[i][1] - Ychastki[i][0])/2 + Ychastki[i][0]) * masshtabX + dx -25, osBalki + 85 - 20, 50, 30, Qt.AlignHCenter, str(round(Ychastki[i][1] - Ychastki[i][0], 3)))

            # -------------------------------------------------------------------------
            ''' Подписи эпюр Mx, Qy, f '''
            qp.setPen(QPen(QtGui.QColor(0, 100, 150), 2, QtCore.Qt.SolidLine))
            qp.setFont(QtGui.QFont("Cambria", 11))
            qp.drawText(dx, 15, 'Расчётная схема балки')
            qp.drawText(dx, osQy - abs(Qmax) * masshtabYQ - 30, 'Эпюра поперечных сил Qy [тс]')
            qp.drawText(dx, osMx - abs(Mmin) * masshtabYM - 30, 'Эпюра изгибающих моментов Mx [тс ∙ м]')
            qp.drawText(dx, osEJvII - abs(EJvIImin) * masshtabYEJvII - 30, 'Эпюра прогибов EIf [тс ∙ m3]')
            
            # ''' Подписи осей эпюр Mx, Qy, f '''
            # qp.setFont(QtGui.QFont("Cambria", 11))
            # qp.drawText(XL[-1] * masshtabX + dx + 10, osQy - 10, 20, 20, Qt.AlignVCenter, 'Qy')
            # qp.drawText(XL[-1] * masshtabX + dx + 10, osMx - 10, 20, 20, Qt.AlignVCenter, 'Mx')
            # qp.drawText(XL[-1] * masshtabX + dx + 10, osEJvII - 10, 20, 20, Qt.AlignVCenter, 'f')

            '''Прорисовка эпюр'''
            for i in range(1, len(XL)):
                ''' Ось Эпюры Qy '''
                qp.setPen(QPen(QtGui.QColor(0, 0, 0), 2, QtCore.Qt.SolidLine))
                qp.drawLine(XL[i-1] * masshtabX + dx, osQy, XL[i] * masshtabX + dx, osQy)
                
                ''' График Эпюры Qy '''
                qp.setPen(QPen(QtGui.QColor(0, 100, 150, 150), 2, QtCore.Qt.SolidLine))
                qp.drawLine(XL[i-1] * masshtabX + dx, osQy - Q_x[i-1] * masshtabYQ, XL[i] * masshtabX + dx, osQy - Q_x[i] * masshtabYQ)

                ''' Ось Эпюры Mx '''
                qp.setPen(QPen(QtGui.QColor(0, 0, 0), 2, QtCore.Qt.SolidLine))
                qp.drawLine(XL[i-1] * masshtabX + dx, osMx, XL[i] * masshtabX + dx, osMx)

                ''' График Эпюры Mx '''
                qp.setPen(QPen(QtGui.QColor(0, 100, 150, 150), 2, QtCore.Qt.SolidLine))
                qp.drawLine(XL[i-1] * masshtabX + dx, osMx + M_x[i-1] * masshtabYM, XL[i] * masshtabX + dx, osMx + M_x[i] * masshtabYM)

                ''' Ось Эпюры EJvII '''
                qp.setPen(QPen(QtGui.QColor(0, 0, 0), 2, QtCore.Qt.SolidLine))
                qp.drawLine(XL[i-1] * masshtabX + dx, osEJvII, XL[i] * masshtabX + dx, osEJvII)

                ''' График Эпюры EJvII '''
                qp.setPen(QPen(QtGui.QColor(0, 100, 150, 150), 2, QtCore.Qt.SolidLine))
                qp.drawLine(XL[i-1] * masshtabX + dx, osEJvII + EJvII[i-1] * masshtabYEJvII, XL[i] * masshtabX + dx, osEJvII + EJvII[i] * masshtabYEJvII)

            ''' Вертикальные линии и значения на границах участкоа '''
            for i in range(0, len(granX)):
                ''' Вертикальные линии на границах участкоа '''
                qp.setPen(QPen(QtGui.QColor(0, 100, 150, 150), 2, QtCore.Qt.SolidLine))
                qp.drawLine(granX[i] * masshtabX + dx, osQy, granX[i] * masshtabX + dx, osQy - granQ[i] * masshtabYQ)
                qp.drawLine(granX[i] * masshtabX + dx, osMx, granX[i] * masshtabX + dx, osMx + granM[i] * masshtabYM)
                qp.setFont(QtGui.QFont("Cambria", 10))
                qp.setPen(QPen(QtGui.QColor(0, 0, 0), 1, QtCore.Qt.SolidLine))

            def dataM(a, b, Align, yyy=0):
                qp.setFont(QtGui.QFont("Cambria", 10))
                qp.setPen(QPen(QtGui.QColor(0, 0, 0), 1, QtCore.Qt.SolidLine))
                if Align == Qt.AlignLeft: xxx = 0
                if Align == Qt.AlignRight: xxx = 70 
                if Align == Qt.AlignHCenter: xxx = 35
                if granMI[a][b] >= 0: yyy = yyy + 3
                if granMI[a][b]  < 0: yyy = - yyy - 17
                qp.drawText(granXI[a][b] * masshtabX + dx - xxx, osMx + granMI[a][b] * masshtabYM + yyy, 70, 30, Align, str(round(granMI[a][b], 2) if granMI[a][b] != 0 else ''))
                
            def dataQ(a, b, Align, yyy=0):
                qp.setFont(QtGui.QFont("Cambria", 10))
                qp.setPen(QPen(QtGui.QColor(0, 0, 0), 1, QtCore.Qt.SolidLine))
                if Align == Qt.AlignLeft: xxx = 0
                if Align == Qt.AlignRight: xxx = 70 
                if Align == Qt.AlignHCenter: xxx = 35
                if granQI[a][b]  < 0: yyy = yyy + 3
                if granQI[a][b] >= 0: yyy = - yyy - 17
                qp.drawText(granXI[a][b] * masshtabX + dx - xxx, osQy - granQI[a][b] * masshtabYQ + yyy, 70, 30, Align, str(round(granQI[a][b], 2) if granQI[a][b] != 0 else ''))
                 
            '''Расстановка значений границ участков эп. Qх'''
            dataQ(0, 0, Qt.AlignLeft)
            dataQ(-1, -1, Qt.AlignRight)
            for i in range(0, len(granQI)-1):
                if granQI[i][-1] == granQI[i+1][0]:
                    dataQ(i, -1, Qt.AlignHCenter)
                else:
                    if granQI[i+1][0] == 0:
                        dataQ(i, -1, Qt.AlignHCenter)
                    if granQI[i][-1] == 0:
                        dataQ(i+1, 0, Qt.AlignHCenter)

                    if granQI[i+1][0] < 0 and granQI[i][-1] > 0:
                        dataQ(i, -1, Qt.AlignHCenter)
                        dataQ(i+1, 0, Qt.AlignHCenter)
                    if granQI[i+1][0] > 0 and granQI[i][-1] < 0:
                        dataQ(i, -1, Qt.AlignHCenter)
                        dataQ(i+1, 0, Qt.AlignHCenter)

                    if granQI[i+1][0] < 0 and granQI[i][-1] < 0:
                        if granQI[i][-1] < granQI[i+1][0]:
                            dataQ(i, -1, Qt.AlignHCenter, 12)
                            dataQ(i+1, 0, Qt.AlignHCenter, abs(abs(granQI[i][-1]) - abs(granQI[i+1][0])) * masshtabYQ - 0)
                        if granQI[i][-1] > granQI[i+1][0]:
                            dataQ(i, -1, Qt.AlignHCenter, abs(abs(granQI[i][-1]) - abs(granQI[i+1][0])) * masshtabYQ - 0)
                            dataQ(i+1, 0, Qt.AlignHCenter, 12)

                    if granQI[i+1][0] > 0 and granQI[i][-1] > 0:
                        if granQI[i][-1] > granQI[i+1][0]:
                            dataQ(i, -1, Qt.AlignHCenter, 12)
                            dataQ(i+1, 0, Qt.AlignHCenter, abs(abs(granQI[i][-1]) - abs(granQI[i+1][0])) * masshtabYQ - 0)
                        if granQI[i][-1] < granQI[i+1][0]:
                            dataQ(i, -1, Qt.AlignHCenter, abs(abs(granQI[i][-1]) - abs(granQI[i+1][0])) * masshtabYQ - 0)
                            dataQ(i+1, 0, Qt.AlignHCenter, 12)


            '''Расстановка значений границ участков эп. Мх'''
            dataM(0, 0, Qt.AlignLeft)
            dataM(-1, -1, Qt.AlignRight)
            for i in range(0, len(granMI)-1):
                if granMI[i][-1] == granMI[i+1][0]:
                    dataM(i, -1, Qt.AlignHCenter)
                else:
                    if granMI[i+1][0] == 0:
                        dataM(i, -1, Qt.AlignHCenter)
                    if granMI[i][-1] == 0:
                        dataM(i+1, 0, Qt.AlignHCenter)

                    if granMI[i+1][0] < 0 and granMI[i][-1] > 0:
                        dataM(i, -1, Qt.AlignHCenter)
                        dataM(i+1, 0, Qt.AlignHCenter)
                    if granMI[i+1][0] > 0 and granMI[i][-1] < 0:
                        dataM(i, -1, Qt.AlignHCenter)
                        dataM(i+1, 0, Qt.AlignHCenter)

                    if granMI[i+1][0] < 0 and granMI[i][-1] < 0:
                        if granMI[i][-1] < granMI[i+1][0]:
                            dataM(i, -1, Qt.AlignHCenter, 12)
                            dataM(i+1, 0, Qt.AlignHCenter, abs(abs(granMI[i][-1]) - abs(granMI[i+1][0])) * masshtabYM - 0)
                        if granMI[i][-1] > granMI[i+1][0]:
                            dataM(i, -1, Qt.AlignHCenter, abs(abs(granMI[i][-1]) - abs(granMI[i+1][0])) * masshtabYM - 0)
                            dataM(i+1, 0, Qt.AlignHCenter, 12)

                    if granMI[i+1][0] > 0 and granMI[i][-1] > 0:
                        if granMI[i][-1] > granMI[i+1][0]:
                            dataM(i, -1, Qt.AlignHCenter, 12)
                            dataM(i+1, 0, Qt.AlignHCenter, abs(abs(granMI[i][-1]) - abs(granMI[i+1][0])) * masshtabYM - 0)
                        if granMI[i][-1] < granMI[i+1][0]:
                            dataM(i, -1, Qt.AlignHCenter, abs(abs(granMI[i][-1]) - abs(granMI[i+1][0])) * masshtabYM - 0)
                            dataM(i+1, 0, Qt.AlignHCenter, 12)

            '''Расстановка максимальных значений между границами участков'''
            for a in range(0, len(MImax)):
                for b in range(0, len(MImax[a])):
                    if MImax[a][b] >= 0: yyy = 3
                    if MImax[a][b]  < 0: yyy = - 17
                    qp.drawText(xImax[a][b] * masshtabX + dx - 35, osMx + MImax[a][b] * masshtabYM + yyy, 70, 30, Qt.AlignHCenter, str(round(MImax[a][b], 2) if MImax[a][b] != 0 else ''))

            ''' Вертикальные линии и максимальные значения в прогибах '''
            def dataV(a, Align, yyy=0):
                qp.setFont(QtGui.QFont("Cambria", 10))
                qp.setPen(QPen(QtGui.QColor(0, 0, 0), 1, QtCore.Qt.SolidLine))
                if Align == Qt.AlignLeft: xxx = 0
                if Align == Qt.AlignRight: xxx = 70 
                if Align == Qt.AlignHCenter: xxx = 35
                if EJvII[a] >= 0: yyy = yyy + 3
                if EJvII[a]  < 0: yyy = - yyy - 17
                qp.drawText(XL[a] * masshtabX + dx - xxx, osEJvII + EJvII[a] * masshtabYEJvII + yyy, 70, 30, Align, str(round(EJvII[a], 2) if EJvII[a] != 0 else ''))
            
            for i in range(0, len(XLEJvII)):
                qp.setPen(QPen(QtGui.QColor(0, 100, 150, 150), 2, QtCore.Qt.SolidLine))
                qp.drawLine(XLEJvII[i] * masshtabX + dx, osEJvII, XLEJvII[i] * masshtabX + dx, osEJvII + EJvII_I[i] * masshtabYEJvII)

            dataV( 0, Qt.AlignLeft)
            dataV(-1, Qt.AlignRight)

            # print('XLEJvII = ', XLEJvII)
            # print('EJvII = ', EJvII)

            for i in range(0, len(XL)):
                if abs(round(EJvII[i], 5)) == abs(EJvIIRmax) or abs(round(EJvII[i], 5)) == abs(EJvIIRmin):
                    dataV(i, Qt.AlignHCenter)

            # ''' Вертикальные линии и максимальные значения в прогибах '''
            # for i in range(0, len(XLEJvII)):
            #     qp.setPen(QPen(QtGui.QColor(0, 100, 150, 150), 2, QtCore.Qt.SolidLine))
            #     qp.drawLine(XLEJvII[i] * masshtabX + dx, osEJvII, XLEJvII[i] * masshtabX + dx, osEJvII + EJvII_I[i] * masshtabYEJvII)
            #     qp.setFont(QtGui.QFont("Cambria", 10))
            #     qp.setPen(QPen(QtGui.QColor(0, 0, 0), 1, QtCore.Qt.SolidLine))
            #     if EJvII_I[i] < 0:
            #         if XLEJvII[i] < L - 0.5:
            #             # qp.drawText(XLEJvII[i] * masshtabX + dx + 3, osEJvII + EJvII_I[i] * masshtabYEJvII - 5, str(EJvII_I[i]))
            #             qp.drawText(XLEJvII[i] * masshtabX + dx - 70, osEJvII + EJvII_I[i] * masshtabYEJvII - 17, 70, 30, Qt.AlignRight, str(EJvII_I[i]))
            #         else:
            #             qp.drawText(XLEJvII[i] * masshtabX + dx - 70, osEJvII + EJvII_I[i] * masshtabYEJvII - 17, 70, 30, Qt.AlignRight, str(EJvII_I[i]))
            #     else:
            #         if XLEJvII[i] < L - 0.5:
            #             qp.drawText(XLEJvII[i] * masshtabX + dx + 3, osEJvII + EJvII_I[i] * masshtabYEJvII + 12, str(EJvII_I[i]))
            #         else:
            #             qp.drawText(XLEJvII[i] * masshtabX + dx -70, osEJvII + EJvII_I[i] * masshtabYEJvII + 5, 70, 30, Qt.AlignRight, str(EJvII_I[i]))

            '''Рисуем вертикальную линейку кратную 5 пикселям для контроля рисунка'''
            # x = 0
            # EEE = []
            # EEE.append(x)
            # while x < Hrisunka:
            #     x += 10
            #     EEE.append(x)
            # # 
            # for i in EEE:
            #     qp.drawLine(0, i, Wrisunka, i)
            #     # qp.drawLine(dx, osBalki + 280, L * masshtabX + dx, osBalki + 280)


    '''====================================================================================================================================='''
    '''====================================================================================================================================='''

    '''Размеры рисунка'''
    Hrisunka = ui.label_28.height()
    Wrisunka = ui.label_28.width()

    '''Запускаем русунок русоваться для проги'''
    ex = Example()

    ee = ex.grab()
    pixmap = QPixmap(ee)
    ui.label_28.setPixmap(pixmap)
    # pixmap.save("imagei.png")

    '''Перерисовываем рисунок для Ворлда'''
    Wrisunka = 600
    Hrisunka = 500
    exxe = Example()
    exxe.setStyleSheet("background-color: rgb(255, 255, 255);")
    ee = exxe.grab()
    pix = QPixmap(ee)
    pix.save("image.png")

    # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
    '''Абсолютное максимальне значение М и Q'''
    MmaxABS = round(max(abs(Mmax), abs(Mmin)), 2)
    QmaxABS = round(max(abs(Qmax), abs(Qmin)), 2)
    # Переводим из тс в Н
    MmaxABS = round(MmaxABS * 9806613.5802, 2)
    QmaxABS = round(QmaxABS * 9806.65, 2)

    '''Модуль упругости стали, Мпа(Н/мм2)'''
    E = dannie[1]  # Мпа(Н/мм2)

    # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
    ui.textEdit_3.setText('')

    '''***********************************************************************************'''
    '''
    Имена значений для двутавра и швеллера
    № - номер двутавра (швеллера)
    h - высота двутавра (швеллера)
    b - ширина полки
    s - толщина стенки
    t - средняя толщина полки
    R - радиус внутреннего закругления
    r - радиус закругления полки
    A - площадь
    m - масса
    Ix (Iy) - момент инерции
    Wx (Wy) - момент сопротивления
    ix (iy) - радиус инерции
    Sx (Sy) - статический момент полусечения
    '''

    if ui.toolBox.currentIndex() == 0:

        '''Двутавры по ГОСТ 8239-89:'''

        DT10   =   ( '10', 100.00, 55.00, 4.50, 7.20, 7.00, 2.50, 12.00, 9.460, 198.000, 39.700, 4.060, 23.000, 17.9000, 6.490, 1.22 )
        DT12   =   ( '12', 120.00, 64.00, 4.80, 7.30, 7.50, 3.00, 14.70, 11.50, 350.000, 58.400, 4.880, 33.700, 27.9000, 8.720, 1.38 )
        DT14   =   ( '14', 140.00, 73.00, 4.90, 7.50, 8.00, 3.00, 17.40, 13.70, 572.000, 81.700, 5.730, 46.800, 41.9000, 11.50, 1.55 )
        DT16   =   ( '16', 160.00, 81.00, 5.00, 7.80, 8.50, 3.50, 20.20, 15.90, 873.000, 109.00, 6.570, 62.300, 58.6000, 14.50, 1.70 )
        DT18   =   ( '18', 180.00, 90.00, 5.10, 8.10, 9.00, 3.50, 23.40, 18.40, 1290.00, 143.00, 7.420, 81.400, 82.6000, 18.40, 1.88 )
        DT20   =   ( '20', 200.00, 100.0, 5.20, 8.40, 9.50, 4.00, 26.80, 21.00, 1840.00, 184.00, 8.280, 104.00, 115.000, 23.10, 2.07 )
        DT22   =   ( '22', 220.00, 110.0, 5.40, 8.70, 10.0, 4.00, 30.60, 24.00, 2550.00, 232.00, 9.130, 131.00, 157.000, 28.60, 2.27 )
        DT24   =   ( '24', 240.00, 115.0, 5.60, 9.50, 10.5, 4.00, 34.80, 27.30, 3460.00, 289.00, 9.970, 163.00, 198.000, 34.50, 2.37 )
        DT27   =   ( '27', 270.00, 125.0, 6.00, 9.80, 11.0, 4.50, 40.20, 31.50, 5010.00, 371.00, 11.20, 210.00, 260.000, 41.50, 2.54 )
        DT30   =   ( '30', 300.00, 135.0, 6.50, 10.2, 12.0, 5.00, 46.50, 36.50, 7080.00, 472.00, 12.30, 268.00, 337.000, 49.90, 2.69 )
        DT33   =   ( '33', 330.00, 140.0, 7.00, 11.2, 13.0, 5.00, 53.80, 42.20, 9840.00, 597.00, 13.50, 339.00, 419.000, 59.90, 2.79 )
        DT36   =   ( '36', 360.00, 145.0, 7.50, 12.3, 14.0, 6.00, 61.90, 48.60, 13380.0, 743.00, 14.70, 423.00, 516.000, 71.10, 2.89 )
        DT40   =   ( '40', 400.00, 155.0, 8.30, 13.0, 15.0, 6.00, 72.60, 57.00, 19062.0, 953.00, 16.20, 545.00, 667.000, 86.10, 3.03 )
        DT45   =   ( '45', 450.00, 160.0, 9.00, 14.2, 16.0, 7.00, 84.70, 66.50, 27696.0, 1231.0, 18.10, 708.00, 808.000, 101.0, 3.09 )
        DT50   =   ( '50', 500.00, 170.0, 10.0, 15.2, 17.0, 7.00, 100.0, 78.50, 39727.0, 1589.0, 19.90, 919.00, 1043.00, 123.0, 3.23 )
        DT55   =   ( '55', 550.00, 180.0, 11.0, 16.5, 18.0, 7.00, 118.0, 92.60, 55962.0, 2035.0, 21.80, 1181.0, 1356.00, 151.0, 3.39 )
        DT60   =   ( '60', 600.00, 190.0, 12.0, 17.8, 20.0, 8.00, 138.0, 108.0, 76806.0, 2560.0, 23.60, 1491.0, 1725.00, 182.0, 3.54 )

        '''Швеллеры по ГОСТ 8240-97'''
        SH5У   =   ( '5У'  , 50.00, 32.00, 4.40, 7.00, 6.00, 2.50, 6.160, 4.840, 22.8000, 9.100, 1.92, 5.590, 5.6100, 2.750, 0.95 )  #, 1.16 - расстояния до центра тяжести
        SH6_5У =   ( '6.5У', 65.00, 36.00, 4.40, 7.20, 6.00, 2.50, 7.510, 5.900, 48.6000, 15.00, 2.54, 9.000, 8.7000, 3.680, 1.08 )  #, 1.24
        SH8У   =   ( '8У'  , 80.00, 40.00, 4.50, 7.40, 6.50, 2.50, 8.980, 7.050, 89.4000, 22.40, 3.16, 13.30, 12.800, 4.750, 1.19 )  #, 1.31
        SH10У  =   ( '10У' , 100.0, 46.00, 4.50, 7.60, 7.00, 3.00, 10.90, 8.590, 174.000, 34.80, 3.99, 20.40, 20.400, 6.460, 1.37 )  #, 1.44
        SH12У  =   ( '12У' , 120.0, 52.00, 4.80, 7.80, 7.50, 3.00, 13.30, 10.40, 304.000, 50.60, 4.78, 29.60, 31.200, 8.520, 1.53 )  #, 1.54
        SH14У  =   ( '14У' , 140.0, 58.00, 4.90, 8.10, 8.00, 3.00, 15.60, 12.30, 491.000, 70.20, 5.60, 40.80, 45.400, 11.00, 1.70 )  #, 1.67
        SH16У  =   ( '16У' , 160.0, 64.00, 5.00, 8.40, 8.50, 3.50, 18.10, 14.20, 747.000, 93.40, 6.42, 54.10, 63.300, 13.80, 1.87 )  #, 1.80
        SH16аУ =   ( '16аУ', 160.0, 68.00, 5.00, 9.00, 8.50, 3.50, 19.50, 15.30, 823.000, 103.0, 6.49, 59.40, 78.800, 16.40, 2.01 )  #, 2.00
        SH18У  =   ( '18У' , 180.0, 70.00, 5.10, 8.70, 9.00, 3.50, 20.70, 16.30, 1090.00, 121.0, 7.24, 69.80, 86.000, 17.00, 2.04 )  #, 1.94
        SH18аУ =   ( '18аУ', 180.0, 74.00, 5.10, 9.30, 9.00, 3.50, 22.20, 17.40, 1190.00, 132.0, 7.32, 76.10, 105.00, 20.00, 2.18 )  #, 2.13
        SH20У  =   ( '20У' , 200.0, 76.00, 5.20, 9.00, 9.50, 4.00, 23.40, 18.40, 1520.00, 152.0, 8.07, 87.80, 113.00, 20.50, 2.20 )  #, 2.07
        SH22У  =   ( '22У' , 220.0, 82.00, 5.40, 9.50, 10.0, 4.00, 26.70, 21.00, 2110.00, 192.0, 8.89, 110.0, 151.00, 25.10, 2.37 )  #, 2.21
        SH24У  =   ( '24У' , 240.0, 90.00, 5.60, 10.0, 10.5, 4.00, 30.60, 24.00, 2900.00, 242.0, 9.73, 139.0, 208.00, 31.60, 2.60 )  #, 2.42
        SH27У  =   ( '27У' , 270.0, 95.00, 6.00, 10.5, 11.0, 4.50, 35.20, 27.70, 4160.00, 308.0, 10.9, 178.0, 262.00, 37.30, 2.73 )  #, 2.47
        SH30У  =   ( '30У' , 300.0, 100.0, 6.50, 11.0, 12.0, 5.00, 40.50, 31.80, 5810.00, 387.0, 12.0, 224.0, 327.00, 43.60, 2.84 )  #, 2.52
        SH33У  =   ( '33У' , 330.0, 105.0, 7.00, 11.7, 13.0, 5.00, 46.50, 36.50, 7980.00, 484.0, 13.1, 281.0, 410.00, 51.80, 2.97 )  #, 2.59
        SH36У  =   ( '36У' , 360.0, 110.0, 7.50, 12.6, 14.0, 6.00, 53.40, 41.90, 10820.0, 601.0, 14.2, 350.0, 513.00, 61.70, 3.10 )  #, 2.68
        SH40У  =   ( '40У' , 400.0, 115.0, 8.00, 13.5, 15.0, 6.00, 61.50, 48.30, 15220.0, 761.0, 15.7, 444.0, 642.00, 73.40, 3.23 )  #, 2.75

        SH5П   =   ( '5П'  , 50.00, 32.00, 4.40, 7.00, 6.00, 3.50, 6.160, 4.84, 22.8000, 9.100, 1.92, 5.6100, 5.950, 2.990, 0.98 )   #, 1.21
        SH6_5П =   ( '6.5П', 65.00, 36.00, 4.40, 7.20, 6.00, 3.50, 7.510, 5.90, 48.8000, 15.00, 2.55, 9.0200, 9.350, 4.060, 1.12 )   #, 1.29
        SH8П   =   ( '8П'  , 80.00, 40.00, 4.50, 7.40, 6.50, 3.50, 8.980, 7.05, 89.8000, 22.50, 3.16, 13.300, 13.90, 5.310, 1.24 )   #, 1.38
        SH10П  =   ( '10П' , 100.0, 46.00, 4.50, 7.60, 7.00, 4.00, 10.90, 8.59, 175.000, 34.90, 3.99, 20.500, 22.60, 7.370, 1.44 )   #, 1.53
        SH12П  =   ( '12П' , 120.0, 52.00, 4.80, 7.80, 7.50, 4.50, 13.30, 10.4, 305.000, 50.80, 4.79, 29.700, 34.90, 9.840, 1.62 )   #, 1.66
        SH14П  =   ( '14П' , 140.0, 58.00, 4.90, 8.10, 8.00, 4.50, 15.60, 12.3, 493.000, 70.40, 5.61, 40.900, 51.50, 12.90, 1.81 )   #, 1.82
        SH16П  =   ( '16П' , 160.0, 64.00, 5.00, 8.40, 8.50, 5.00, 18.10, 14.2, 750.000, 93.80, 6.44, 54.300, 72.80, 16.40, 2.00 )   #, 1.97
        SH16аП =   ( '16аП', 160.0, 68.00, 5.00, 9.00, 8.50, 5.00, 19.50, 15.3, 827.000, 103.0, 6.51, 59.500, 90.50, 19.60, 2.15 )   #, 2.19
        SH18П  =   ( '18П' , 180.0, 70.00, 5.10, 8.70, 9.00, 5.00, 20.70, 16.3, 1090.00, 121.0, 7.26, 70.000, 100.0, 20.60, 2.20 )   #, 2.14
        SH18аП =   ( '18аП', 180.0, 74.00, 5.10, 9.30, 9.00, 5.00, 22.20, 17.4, 1200.00, 133.0, 7.34, 76.300, 123.0, 24.30, 2.35 )   #, 2.36
        SH20П  =   ( '20П' , 200.0, 76.00, 5.20, 9.00, 9.50, 5.50, 23.40, 18.4, 1530.00, 153.0, 8.08, 88.000, 134.0, 25.20, 2.39 )   #, 2.30
        SH22П  =   ( '22П' , 220.0, 82.00, 5.40, 9.50, 10.0, 6.00, 26.70, 21.0, 2120.00, 193.0, 8.90, 111.00, 178.0, 31.00, 2.58 )   #, 2.47
        SH24П  =   ( '24П' , 240.0, 90.00, 5.60, 10.0, 10.5, 6.00, 30.60, 24.0, 2910.00, 243.0, 9.75, 139.00, 248.0, 39.50, 2.85 )   #, 2.72
        SH27П  =   ( '27П' , 270.0, 95.00, 6.00, 10.5, 11.0, 6.50, 35.20, 27.7, 4180.00, 310.0, 10.9, 178.00, 314.0, 46.70, 2.99 )   #, 2.78
        SH30П  =   ( '30П' , 300.0, 100.0, 6.50, 11.0, 12.0, 7.00, 40.50, 31.8, 5830.00, 389.0, 12.0, 224.00, 393.0, 54.80, 3.12 )   #, 2.83
        SH33П  =   ( '33П' , 330.0, 105.0, 7.00, 11.7, 13.0, 7.50, 46.50, 36.5, 8010.00, 486.0, 13.1, 281.00, 491.0, 64.60, 3.25 )   #, 2.90
        SH36П  =   ( '36П' , 360.0, 110.0, 7.50, 12.6, 14.0, 8.50, 53.40, 41.9, 10850.0, 603.0, 14.3, 350.00, 611.0, 76.30, 3.38 )   #, 2.99
        SH40П  =   ( '40П' , 400.0, 115.0, 8.00, 13.5, 15.0, 9.00, 61.50, 48.3, 15260.0, 763.0, 15.8, 445.00, 760.0, 89.90, 3.51 )   #, 3.05

        DTALL = [DT10, DT12, DT14, DT16, DT18, DT20, DT22, DT24, DT27, DT30, DT33, DT36, DT40, DT45, DT50, DT55, DT60]
        SHУALL = [SH5У, SH6_5У, SH8У, SH10У, SH12У, SH14У, SH16У, SH16аУ, SH18У, SH18аУ, SH20У, SH22У, SH24У, SH27У, SH30У, SH33У, SH36У, SH40У]
        SHПALL = [SH5П, SH6_5П, SH8П, SH10П, SH12П, SH14П, SH16П, SH16аП, SH18П, SH18аП, SH20П, SH22П, SH24П, SH27П, SH30П, SH33П, SH36П, SH40П]

        '''--------------------------------------------------------------------------------'''
        
        if ui.comboBox_1.currentText() == ProfALL[0]:       # Если двутавр
            DataEl = DTALL[ui.comboBox_2.currentIndex()]
        if ui.comboBox_1.currentText() == ProfALL[1]:       # Если швеллер У
            DataEl = SHУALL[ui.comboBox_2.currentIndex()]
        if ui.comboBox_1.currentText() == ProfALL[2]:       # Если швеллер П
            DataEl = SHПALL[ui.comboBox_2.currentIndex()]

        DTSHName = ('№', 'h', 'b', 's', 't', 'R', 'r', 'A', 'm', 'Ix', 'Wx', 'ix', 'Sx', 'Iy', 'Wy', 'iy')
        element = {DTSHName[i] : DataEl[i] for i in range(0, len(DTSHName))}

        b = element['b']
        tw = element['t']
        h2w = round(element['h'] - element['t'], 1)
        h = element['h']
        Wx = element['Wx']
        Wx = round(Wx * 1000)
        Sx = element['Sx'] * 1000
        Ix = element['Ix'] * 10000
        Iy = round(element['Iy'] * 10000)
        A = element['A']

        '''Расчет площадей полок, стенок и их отношений друг к другу двутавра'''
        # Площадь полки
        Afx = element['t'] * element['b']
        # Площадь стенки с вычетом средних толщин двух полок
        Awx = element['s'] * (element['h'] - 2 * element['t'])
        '''Отношение площади полки к площади стенки при изгибе вокруг оси X afwx = 0,88'''
        afw = round(Afx / Awx, 2)

        Aw = A/(1 + 2 * afw)
        Af = A * (afw/(1 + 2 * afw)) 
        # Переводим из см2 в мм2
        Aw = round(Aw * 100, 4)
        Af = round(Af * 100, 4)

    if ui.toolBox.currentIndex() == 1:
        RtT = vvod('_29', 2, 0)
        D = RtT[0]
        d = D - 2 * s
        s = RtT[1]
        R = RtT[0] / 2
        r = (RtT[0] - 2 * s) / 2
        c = (2 * r) / (2 * R)
        if s <= D / 10:
            rsr = R - 0.5 * s
            dsr = 2 * rsr
            do = dsr
            Ix = pi * rsr**3 * s
            Wx = pi * rsr**2 * s
        else:
            Ix = (pi * rsr**4 / 4) * (1 - c**4)
            Wx = (pi * D**3 / 32) * (1 - c**4)
        Yo = (4/(3 * pi)) * ((R**2 + R * r + r**2) / (R + r))
        F =  pi * (R**2 - r**2) / 2
        Sx = F * Yo
        tw = s
        Ix = round(Ix, 2)
        Wx = round(Wx, 2)
        Sx = round(Sx, 2)

    if ui.toolBox.currentIndex() == 2 or ui.toolBox.currentIndex() == 3:
        '''Выборр списка квадратное или прямоугольное'''
        if ui.toolBox.currentIndex() == 2:
            indexx = ui.comboBox_6.currentIndex()
            elem = KvadList[indexx]
        if ui.toolBox.currentIndex() == 3:
            indexx = ui.comboBox_7.currentIndex()
            elem = PraymoygList[indexx]
        '''Сбор характеристик'''
        Ix = elem['Ix']
        Wx = elem['Wx']
        h  = elem['h']
        b  = elem['b']
        tw = elem['t']
        A  = elem['A']
        if ui.toolBox.currentIndex() == 2:
            Iy = Ix
            Wy = Wx
        if ui.toolBox.currentIndex() == 3:
            Iy = elem['Iy']
            Wy = elem['Wy']
        '''Вычисляем малый радиус - r '''
        if tw <= 6: R = 2 * tw
        if 6 < tw <= 10: R = 2.5 * tw
        if tw > 10: R = 3 * tw
        r = R - tw
        '''Площадь закругления'''
        FR = pi * 90 / 360 * (R**2 - r**2)
        '''Площадь полки без закругления'''
        Fpol = (b - R * 2) * tw
        '''Площадь стенки без закругления'''
        Fste = (h - R) * tw
        '''Центр тяжести закругления относительно локальной Ц.Т.'''
        YR = (4 / 3) * ((R**3 - r**3) / (R**2 - r**2)) * (180 / (pi * 90))*sin(90 / 2)
        '''Центр тяжести закругления относительно глобального Ц.Т.'''
        YoR = (h / 2 - tw) + ((YR**2 / 2)**0.5)
        '''Центр тяжести полки относительно глобального Ц.Т.'''
        Ypol = (h - tw/2)
        '''Центр тяжести стенки относительно глобального Ц.Т.'''
        Yste = (h - R) * 0.5
        '''Статический момент полусечения = сумме всех площадей на глобальные Ц.Т.'''
        Sx = (FR * YoR * 2 + Fpol * Ypol + Fste * Yste * 2)
        # Момент инерции сечения без зауруглений
        # Ix = ((b * h**3) - (b - 2 * tw) * (h - 2 * tw)**3) / 12

        '''Переводим из см2 в мм2'''
        Ix = round(Ix * 10000)
        Iy = round(Iy * 10000)
        Wx = round(Wx * 1000)
        Wy = round(Wy * 1000)


        '''Расчет площадей полок, стенок и их отношений друг к другу двутавра'''
        # Площадь полки с учетом двух половинок закруглений
        Afx = (Fpol + FR * 2) * 2
        # Площадь стенки с учетом двух половинок закруглений
        Awx = (Fste + FR * 2) * 2
        '''Отношение площади полки к площади стенки при изгибе вокруг оси X afwx = 0,88'''
        afw = round(Afx / Awx, 2)

        Aw = A/(1 + 2 * afw)
        Af = A * (afw/(1 + 2 * afw)) 
        # Переводим из см2 в мм2
        Aw = round(Aw * 100, 4)
        Af = round(Af * 100, 4)
        h2w = round(h - tw, 1)

    if ui.toolBox.currentIndex() == 4:
        SosDv = vvod('_33', 6, 0)
        print('SosDv = ', SosDv)
        '''Высота сечения'''
        h = SosDv[0]
        '''Толщина стенки'''
        tw = SosDv[1]
        '''Ширина верхней полки'''
        bvp = SosDv[2]
        '''Толщтна верхней полки'''
        tvp = SosDv[3]
        '''Ширина нижней полки'''
        bnp = SosDv[4]
        '''Толщтна нижней полки'''
        tnp = SosDv[5]
        '''Высота стенки'''
        hw = h - tvp - tnp
        '''Площадь верхней полки'''
        Avp = bvp * tvp
        '''Площадь нижней полки'''
        Anp = bnp * tnp
        '''Площадь стенки полки'''
        Ast = (h - (tvp + bnp)) * tw
        '''Площадь сечения'''
        A = Avp + Anp + Ast
        h2w = round(h - (tvp + bnp) * 0.5, 1)
        afw = round(Avp / Ast, 2)
        Aw = A/(1 + 2 * afw)
        Af = A * (afw/(1 + 2 * afw)) 
        # Переводим из см2 в мм2
        Aw = round(Aw * 100, 4)
        Af = round(Af * 100, 4)

        '''     Для определения центра тяжести составного сечения выбираем исход-ные оси x и y. 
        Эти оси желательно выбрать так, чтобы элементы сечения находились в первом квадранте координатных осей. 
        Через центры тяжести каждого элемента проводим собственные центральные оси x1 и y1, x2 и y2, x3 и y3, параллельные осям x и y.
        Самая левая и нижняя точка сечения принята нулевой координатой XOY.  '''
        
        '''Ц.Т. относительно XOY'''
        Yc_vp = h - tvp * 0.5
        Yc_np = tnp * 0.5
        Yc_st = tnp + hw * 0.5
        '''Координаты Ц.Т. полок и стенки относительно XOY'''
        Yc = (Avp * Yc_vp + Anp * Yc_np + Ast * Yc_st) / A
        # Xc = max(bvp, bnp) / 2
        '''Осевой момент инерции X верхней полки, стенки и ниж. полки'''
        Ix_vp = bvp * tvp**3 / 12
        Ix_np = bnp * tnp**3 / 12
        Ix_st = tw * hw**3 / 12
        '''Осевой момент инерции Y верхней полки, стенки и ниж. полки'''
        Iy_vp = bvp**3 * tvp / 12
        Iy_np = bnp**3 * tnp / 12
        Iy_st = tw**3 * hw / 12
        '''Расстояние от Ц.Т. сечения до Ц.Т. каждой элемента сечения'''
        Yco_vp = Yc - Yc_vp
        Yco_np = Yc - Yc_np
        Yco_st = Yc - Yc_st
        '''Определим моменты инерции всего сечения относительно центральных осей'''
        Ix = (Ix_vp + Yco_vp**2 * Avp) + (Ix_np + Yco_np**2 * Anp) + (Ix_st + Yco_st**2 * Ast)
        Ix = round(Ix, 2)
        Iy = Iy_vp + Iy_np + Iy_st
        Iy = round(Iy, 2)
        '''Момент сопротивления'''
        Wx = round(Ix / (max(h - Yc, Yc)), 2)
        Wy = round(Iy / (max(bvp / 2, bnp / 2)), )
        '''Ц.Т. полусечения'''
        Yc_polyStenki = (h - Yc - tvp) / 2
        # Yc_polySPolki = (h - Yc - tvp * 0.5)
        '''Статический момент инерциии полусечения'''
        Sx = round(Avp * abs(Yco_vp) + ((Ast * 0.5) * Yc_polyStenki), 2)



    # Функция интерполяции 
    def interpoi (t1, t2, yy):
        for i in t1:
            if yy > i:
                continue
            else:
                break
        ia1 = t1.index(i); ia2 = t1.index(i)-1
        a1 = t1[ia1]; a2 = t1[ia2]; b1 = t2[ia1]; b2 = t2[ia2]
        a12 = a1 - a2; b12 = b1 - b2; a13 = yy - a2; x = a13 * b12 / a12; y = b2 + x
        return y

    '''Таблица Е.1'''
    # '''Двутавр'''
    tDT = (0.25, 0.5, 1.0, 2.0)
    tDTCx = (1.19, 1.12, 1.07, 1.04)
    # '''Швеллер'''
    tSH = (0.5, 1.0, 2.0)
    tSHCx = (1.07, 1.12, 1.19)
    # '''Труба'''
    tTRUBACx = 1.26
    
    ''' 
    ProfALL[0] = Двутавр ГОСТ 8239-89 
    ProfALL[1] = Швеллер У ГОСТ 8240-97
    ProfALL[2] = Швеллер П ГОСТ 8240-97
    ProfALL[3] = Труба ГОСТ 10704-91
    '''

    TipSech = ui.comboBox_1.currentText()
    # Двутавр
    if TipSech == ProfALL[0] and ui.toolBox.currentIndex() == 0:
        Сx = round(interpoi (tDT, tDTCx, afw), 4)
    if ui.toolBox.currentIndex() != 1:
        Сx = round(interpoi (tDT, tDTCx, afw), 4)
    # Швеллер
    if ui.toolBox.currentIndex() == 0:
        if TipSech == ProfALL[1] or TipSech == ProfALL[2]:
            Сx = round(interpoi (tSH, tSHCx, afw), 4)
    # Труба
    if ui.toolBox.currentIndex() == 1:
        Сx = tTRUBACx

    Cxm = round(0.5 * (1 + Сx), 4)
    # Коэффициент условия работы
    yc = dannie[2]

    NDS = ui.comboBox_0.currentIndex()

    '''======================================================================================='''

    '''======================================================================================='''

    # text_centr('Расчет элементов стальных конструкций при изгибе по СП 16.13330.2017')

    paragraph = document.add_paragraph()
    # paragraph.add_run('Геологический разрез', style='Intense Emphasis').bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    '''Добавляем рисунок в документ для сохранения в Ворлд'''
    document.add_picture('image.png', width=Cm(17.0))
    # document.add_picture('image.png', height=Cm(15.0))
    '''Выравнивание русунка в ворлде по центру'''
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    '''Продолжаем собирарть отчет'''
    if TipSech == ProfALL[0] and ui.toolBox.currentIndex() == 0:
        x = 'двутавр'
        z = '8239-89'
    if TipSech == ProfALL[1] or TipSech == ProfALL[2] and ui.toolBox.currentIndex() == 0:
        x = 'швеллер'
        z = '8240-97'
    if ui.toolBox.currentIndex() == 1:
        x = 'труба'
        z = '10704-91'
        DataEl = [str(D) + " x " +  str(s)]
    if ui.toolBox.currentIndex() == 2:
        x = 'профиль гнутый замкнутый квадратный'
        z = '30245-2003 '
        DataEl = KvadListName
    if ui.toolBox.currentIndex() == 3:
        x = 'профиль гнутый замкнутый прямоугольный'
        z = '30245-2003 '
        DataEl = PraymoygListName

    if ui.toolBox.currentIndex() == 0:
        ui.textEdit_3.append('')
        text_abzac('Сечение балки  -  {} № {} по ГОСТ {}, марка стали {}.'.format(x, DataEl[0], z, MarkaStali))

    if ui.toolBox.currentIndex() == 2:
        ui.textEdit_3.append('')
        text_abzac('Сечение балки  -  {} № {} по ГОСТ {}, марка стали {}.'.format(x, DataEl[ui.comboBox_6.currentIndex()], z, MarkaStali))

    if ui.toolBox.currentIndex() == 3:
        ui.textEdit_3.append('')
        text_abzac('Сечение балки  -  {} № {} по ГОСТ {}, марка стали {}.'.format(x, DataEl[ui.comboBox_7.currentIndex()], z, MarkaStali))

    if ui.toolBox.currentIndex() == 1:
        ui.textEdit_3.append('')
        text_abzac('Сечение балки  -  {} {} мм по ГОСТ {}, марка стали {}.'.format(x, DataEl[0], z, MarkaStali))

    if ui.toolBox.currentIndex() == 4:
        ui.textEdit_3.append('')
        text_abzac('Сечение балки  -  составной двутавр с размерами сечения:')
        ui.textEdit_3.append('')
        text_abzac_left('''        высота профиля - {} мм, 
        толщина стенки - {} мм, 
        ширина верхней полки - {} мм,
        толщина верхней полки - {} мм,
        ширина нижней полки - {} мм,
        толщина нижней полки - {} мм.'''.format(h, tw, bvp, tvp, bnp, tnp))
        b = bvp
    # text_abzac('')

    ui.textEdit_3.setFontUnderline(True)
    text_centr('Результаты расчета:')
    ui.textEdit_3.setFontUnderline(False)
    ui.textEdit_3.append('')

    if NDS == 1:    # если упругое-пластичное состояние балки
        text_abzac('8.2.3    Расчет на прочность разрезных балок из стали с нормативным сопротивлением Ry ≤ 440 МПа , несущих статическую нагрузку, с учетом развития пластических деформаций, следует выполнять по формулам: ')
        ui.textEdit_3.append('')
        text_abzac('Расчет на прочность не в опорном сечении балок (при Qy = 0):')
        text_centr('Mx / (Сxm ∙ β ∙ Wxn,min ∙ Ry ∙ γc) ≤ 1')
        
        G = round(MmaxABS / (Cxm * Wx * Ry * yc), 4)
        if G == 0:
            error_show('Укажите нагрузку на балку')
            return
        print('G = ', G)
        text_abzac('где:')
        text_abzac('Mx = {} H мм  -  абсолютное значения изгибающего момента;'.format(MmaxABS))
        text_abzac('Сxm = 0.5 ∙ (1 + Сx) = 0.5 ∙ (1 + {}) = {} - где, Cx = {}  -  коэффициент, принимаемый по таблице Е.1;'.format(Сx, Cxm, Сx))
        text_abzac('β = 1  -  коэффициент в зоне чистого изгиба;')
        text_abzac('Wxn,min = {} мм3 -  момент сопротивления сечения;'.format(Wx))
        text_abzac('Ry = {} Н/мм2  -  расчетное сопротивление растяжению, сжатию, изгибу по пределу текучести;'.format(Ry))
        text_abzac('yc = {}  -  коэффициент условий работы.'.format(yc))
        if G <= 1:
            text_centr('{} Н мм / ({} ∙ 1 ∙ {} мм3 ∙ {} Н/мм2 ∙ {}) = {}  ≤  1'.format(MmaxABS, Cxm, Wx, Ry, yc, G))
        else:
            text_centr('{} Н мм / ({} ∙ 1 ∙ {} мм3 ∙ {} Н/мм2 ∙ {}) = {}  >  1'.format(MmaxABS, Cxm, Wx, Ry, yc, G))
            text_centr_red('Прочность НЕ обеспечена  -  необходимо увеличить сечение')
            return
        ui.textEdit_3.append('')

        if ui.toolBox.currentIndex() != 1:
            '''Расчет на прочность в опорном сечении балок (при Mx = 0)'''
            '''Находим значения Мх на опорах'''
            Mx_xR1 = abs(M_x[XL.index(xR1)])
            Mx_xR2 = abs(M_x[XL.index(xR2)])
            if Mx_xR1 + Mx_xR2 == 0:
                text_abzac('Расчет на прочность в опорном сечении балок (при Mx = 0):')
                text_centr('Qy / (2 ∙ Af ∙ Rs ∙ γc) ≤ 1')

                QMx0 = round(QmaxABS / (2 * Af * Rs * yc), 4)

                text_abzac('где:')
                text_abzac('Qy = {} H  -  абсолютное значения поперечной силы;'.format(QmaxABS))
                text_abzac('Af = {} Н/мм2  -  площадь полки профиля;'.format(Af))
                text_abzac('Rs = {} Н/мм2  -  расчетное сопротивление стали сдвигу;'.format(Rs))
                text_abzac('yc = {}  -  коэффициент условий работы;'.format(yc))
                text_abzac('β = 1  -  коэффициент в зоне чистого изгиба.')
                text_centr('{} Н / (2 ∙ {} мм2 ∙ {} Н/мм2 ∙ {}) = {}  ≤  1'.format(QmaxABS, Af, Rs, yc, QMx0))
                if QMx0 > 1:
                    text_centr_red('Прочность НЕ обеспечена  -  необходимо увеличить сечение')
                    return
                ui.textEdit_3.append('')

        '''Следующая формула для касательных напряжений кольца сходитсся с общей формулой'''
        if ui.toolBox.currentIndex() == 1:
            text_abzac('при действии в сечении поперечных (касательных) сил по формуле:')
            text_centr('2 ∙ Qy / (3.1416 ∙ do ∙ t ∙ Rs ∙ yc)  ≤  1')
            t = round(tw * 0.5, 2)
            Tay = round((2 * QmaxABS) / (pi * do * t * Rs * yc), 4)
            print('Tay = ', Tay)
            text_abzac('где:')
            text_abzac('Qy = {} H  -  абсолютное значения изгибающего момента;'.format(QmaxABS))
            text_abzac('do = {} мм  -  средний диаметр;'.format(do))
            text_abzac('t = {} мм4  -  половина толщины стенки;'.format(t))
            text_abzac('Rs = {} Н/мм2  -  расчетное сопротивление стали сдвигу;'.format(Rs))
            text_abzac('yc = {}  -  коэффициент условий работы.'.format(yc))
            text_centr('2 ∙ {} H / (3.1416 ∙ {} мм ∙ {} мм ∙ {} Н/мм2 ∙ {}) = {} ≤  1'.format(QmaxABS, do, t, Rs, yc, Tay))
            if Tay > 1:
                text_centr_red('Прочность НЕ обеспечена  -  необходимо увеличить сечение')
                return
            ui.textEdit_3.append('')

    if NDS == 0:    # если упругое состояние балки
        text_abzac('8.2.1    Расчет на прочность балок 1-го класса следует выполнять по формулам:')
        ui.textEdit_3.append('')
        text_abzac('при действии момента в одной из главных плоскостей:')
        text_centr('Mx / (Wxn,min ∙ Ry ∙ γc) ≤ 1')

        G = round(MmaxABS / (Wx * Ry * yc), 4)
        if G == 0:
            error_show('Укажите нагрузку на балку')
            return
        print('G = ', G)
        text_abzac('где:')
        text_abzac('Mx = {} H мм  -  абсолютное значения изгибающего момента;'.format(MmaxABS))
        text_abzac('Wxn,min = {} мм3 -  момент сопротивления сечения;'.format(Wx))
        text_abzac('Ry = {} Н/мм2  -  расчетное сопротивление растяжению, сжатию, изгибу по пределу текучести;'.format(Ry))
        text_abzac('yc = {}  -  коэффициент условий работы.'.format(yc))
        if G > 1:
            text_centr('{} Н мм / ({} мм3 ∙ {} Н/мм2 ∙ {}) = {}  >  1'.format(MmaxABS, Wx, Ry, yc, G))
            text_centr_red('Прочность НЕ обеспечена  -  необходимо увеличить сечение')
            return
        else:
            text_centr('{} Н мм / ({} мм3 ∙ {} Н/мм2 ∙ {}) = {}  ≤  1'.format(MmaxABS, Wx, Ry, yc, G))
        ui.textEdit_3.append('')

        text_abzac('при действии в сечении поперечных (касательных) сил по формуле:')
        text_centr('Qy ∙ Sx / (Ix ∙ tw ∙ Rs ∙ yc)  ≤  1')
        Tay = round((QmaxABS * Sx) / (Ix * tw * Rs * yc), 4)
        print('Tay = ', Tay)
        text_abzac('где:')
        text_abzac('Qy = {} H  -  абсолютное значения изгибающего момента;'.format(QmaxABS))
        text_abzac('Sx = {} мм3  -  статический момент полусечения;'.format(Sx))
        text_abzac('Ix = {} мм4  -  момент инерции сечания;'.format(Ix))
        text_abzac('tw = {} мм  -  толщина стенки сечания;'.format(tw))
        text_abzac('Rs = {} Н/мм2  -  расчетное сопротивление стали сдвигу;'.format(Rs))
        text_abzac('yc = {}  -  коэффициент условий работы.'.format(yc))
        text_centr('{} H ∙ {} мм3 / ({} мм4 ∙ {} мм ∙ {} Н/мм2 ∙ {}) = {} ≤  1'.format(QmaxABS, Sx, Ix, tw, Rs, yc, Tay))
        if Tay > 1:
            text_centr_red('Прочность НЕ обеспечена  -  необходимо увеличить сечение')
            return
        ui.textEdit_3.append('')
        
        '''Следующая формула для касательных напряжений кольца сходитсся с общей формулой'''
        # if ui.toolBox.currentIndex() == 1:
        #     text_abzac('при действии в сечении поперечных (касательных) сил по формуле:')
        #     text_centr('2 ∙ Qy / (3.1416 ∙ do ∙ t ∙ Rs ∙ yc)  ≤  1')
        #     t = round(tw * 0.5, 2)
        #     Tay = round((2 * QmaxABS) / (pi * do * t * Rs * yc), 4)
        #     print('Tay = ', Tay)

        #     text_abzac('где:')
        #     text_abzac('Qy = {} H  -  абсолютное значения изгибающего момента;'.format(QmaxABS))
        #     text_abzac('do = {} мм  -  средний диаметр;'.format(do))
        #     text_abzac('t = {} мм4  -  половина толщины стенки;'.format(t))
        #     text_abzac('Rs = {} Н/мм2  -  расчетное сопротивление стали сдвигу;'.format(Rs))
        #     text_abzac('yc = {}  -  коэффициент условий работы.'.format(yc))
        #     text_centr('2 ∙ {} H / (3.1416 ∙ {} мм ∙ {} мм ∙ {} Н/мм2 ∙ {}) = {} ≤  1'.format(QmaxABS, do, t, Rs, yc, Tay))
        #     if Tay > 1:
        #         text_centr_red('Прочность НЕ обеспечена  -  необходимо увеличить сечение')
        #         return
        #     ui.textEdit_3.append('')

    Stop841 = 'No'
    '''Если не труба круглая, то считаем устойчивость'''
    if ui.toolBox.currentIndex() != 1:
    # if ui.toolBox.currentIndex() == 0 or ui.toolBox.currentIndex() == 4:
        text_abzac('8.4    Расчет на общую устойчивость изгибаемых элементов сплошного сечения')
        ui.textEdit_3.append('')
        
        text_abzac('8.4.4    Устойчивость балок 1-го класса, а также бистальных балок 2-го класса следует считать обеспеченной:')
        ui.textEdit_3.append('')
        text_abzac('а) при передаче нагрузки на балку через сплошной жесткий настил - не выполнено;')
        # text_abzac_insert_color('не выполнено')
        # ui.textEdit_3.insertPlainText(';')
        text_abzac('б) при значении условной гибкости сжатого пояса балки, не превышающего ее предельного значения:')
        ui.textEdit_3.append('')
        text_abzac('Условная предельная гибкость сжатого пояса балки определяется согласно таблице 11:')
        text_centr('λab = 0.35 + 0.0032 ∙ b/t + (0.76 - 0.02 ∙ b/t) ∙ b/h'.format())
        ui.textEdit_3.append('')
        text_abzac('при следующих условиях:')
        text_centr('1 ≤ h/b ≤ 6 и 15 ≤ b/t ≤ 35')
        text_abzac('где:')
        text_abzac('b = {} мм и t = {} мм  -  соответственно ширина и толщина сжатого пояса;'.format(b, tw))
        text_abzac('h = {} мм  -  расстояние (высота) между осями поясных листов.'.format(h2w))
        text_abzac('Для балок с отношением b/t < 15 в формулах таблицы 11 следует принимать b/t = 15.')
    
        hb = round(h2w/b, 4)
        b_h = round(b/h2w, 4)
        b_t = 15 if round(b/tw, 4) < 15 else round(b/tw, 4)
        if 1 <= hb <= 6 and 15 <= b_t <= 35:
            text_centr('1 ≤ {} ≤ 6 и 15 ≤ {} ≤ 35  -  условия выполняются'.format(round(hb, 4), round(b_t, 4)))
            '''Место приложения нагрузки'''
            mesto = {
                    'К верхнему поясу' : 0.35 + 0.0032 * b/tw + (0.76 - 0.02 * b/tw) * b/h2w, 
                    'К нижнему поясу' : 0.57 + 0.0032 * b/tw + (0.92 - 0.02 * b/tw) * b/h2w, 
                    'Независимо от уровня приложения нагрузки или при чистом изгибе' : 0.41 + 0.0032 * b/tw + (0.73 - 0.016 * b/tw) * b/h2w
                    }
            lyab = round(mesto['К верхнему поясу'], 4)
            print('lyab = ', lyab)
            text_centr('λab = 0.35 + 0.0032 ∙ {} + (0.76 - 0.02 ∙ {}) ∙ {} = {}'.format(b_t, b_t, b_h, lyab))
        
            ui.textEdit_3.append('')
            text_abzac('Условная гибкость сжатого пояса:')
            text_centr('λb = (lef / b) ∙ (Ryf / E)^0.5')
            text_abzac('где:')
    
            lef = round(sum(Lbk) * 1000)
            
            text_abzac('lef = {} мм  -  расстояние между точками закреплений сжатого пояса от поперечных смещений;'.format(lef))
            text_abzac('b = {} мм  -  ширина сжатого пояса;'.format(b))
            text_abzac('Ryf = Ry = {} Н/мм2  -  расчетное сопротивление стали сжатого пояса;'.format(Ry))
            text_abzac('E = {} Н/мм2  -  модуль упругости стали;'.format(E))
            lyb = round((lef / b) * (Ry / E)**0.5, 4)
            text_centr('λb = ({} мм / {} мм) ∙ ({} Н/мм2 / {} Н/мм2)^0.5 = {}'.format(lef, b, Ry, E, lyb))
            print('lyb = ', lyb)

            if NDS == 1:      # 'упруго-пластическое состояние сечения':
                ui.textEdit_3.append('')
                text_abzac('''8.4.6    Устойчивость балок 2-го и 3-го классов следует считать обеспеченной при выполнении требований перечисления а) или б) 8.4.4 при условии умножения значений λab, определяемых по формулам таблицы 11, на коэффициент δ:''')
                text_centr('δ = 1 - 0.6 ∙ (C1x - 1) / (Cx-1)')
    
                C1x_1 = round(MmaxABS / (Wx * Ry * yc), 4)
                C1x_2 = 1 * Сx
                C1x = max(C1x_1, C1x_2)
                sigma = 1 - 0.6 * (C1x - 1) / (Сx - 1)
    
                text_abzac('где:')
                text_abzac('Cx1 = {}  -  коэффициент, определяемый по большему значению из формул:'.format(C1x))
                text_abzac('C1x = Mx / (Wx ∙ Ry ∙ yc) = {} Н мм / ({} мм3 ∙ {} Н/мм2 ∙ {}) = {};'.format(MmaxABS, Wx, Ry, yc, C1x_1))
                text_abzac('C1x = β ∙ yc = 1 ∙ {} = {}'.format(yc, C1x_2)) 
                text_abzac('Cx = {}  -  коэффициент, принимаемый по таблице Е.1;'.format(Сx))
                text_abzac('β = 1  -  коэффициент в зоне чистого изгиба, согласно 8.2.3.')
                text_centr('δ = 1 - 0.6 ∙ ({} - 1) / ({} - 1) = {}'.format(C1x, Сx, sigma))
            
            if lyb >= lyab:
                if NDS == 0:    # 'упругое состояние сечения':
                    text_centr('λb = {} ≥ λab = {}  -  условие НЕ выполняется'.format(lyb, lyab))
                    print('λb = {} ≥ λab = {}  -  условие НЕ выполняется'.format(lyb, lyab))
                    ui.textEdit_3.append('')
                else:
                    text_centr('λb = {} ≥ δ ∙ λab = {} ∙ {} = {}  -  условие НЕ выполняется'.format(lyb, sigma, lyab, sigma * lyab))
                    ui.textEdit_3.append('')
                Stop841 = 'No'
                text_abzac('Так как условие НЕ выполняется, требуется проверить устойчивость балки по 8.4.1.')
            else:
                if NDS == 0:    # 'упругое состояние сечения':
                    text_centr('λb = {} ≤ λab = {}  -  условие выполняется'.format(lyb, lyab))
                    print('λb = {} ≤ λab = {}  -  условие выполняется'.format(lyb, lyab))
                else:
                    text_centr('λb = {} ≤ δ ∙ λab = {} ∙ {} = {}  -  условие выполняется'.format(lyb, sigma, lyab, sigma * lyab))
                # ui.textEdit_3.append('')
                text_centr('Устойчивость балки обеспечена.')
                Stop841 = 'Yes'
                print('Устойчивость балки обеспечена')
                print('Stop841 = ', Stop841)
                # return
        else:
            ui.textEdit_3.setTextColor(QtGui.QColor (255, 0, 0))
            text_centr('1 ≤ {} ≤ 6 и 15 ≤ {} ≤ 35  -  условия НЕ выполняется'.format(round(hb, 4), round(b_t, 4)))
            ui.textEdit_3.setTextColor(QtGui.QColor (0, 0, 0))
            ui.textEdit_3.append('')
            text_abzac('Так как условия НЕ выполняется, требуется проверить устойчивость балки по 8.4.1')
            Stop841 = 'No'
            print('1 ≤ {} ≤ 6 и 15 ≤ {} ≤ 35  -  условия НЕ выполняется'.format(round(hb, 4), round(b_t, 4)))
            print('Stop841 = ', Stop841)
        ui.textEdit_3.append('')



        if Stop841 != 'Yes':
            if TipSech == ProfALL[0] and ui.toolBox.currentIndex() == 0: kI = 1.29
            if ui.toolBox.currentIndex() == 4: kI = 1.29
            if TipSech == ProfALL[1] and ui.toolBox.currentIndex() == 0: kI = 1.12
            if TipSech == ProfALL[2] and ui.toolBox.currentIndex() == 0: kI = 1.12
            if ui.toolBox.currentIndex() == 2 or ui.toolBox.currentIndex() == 3: kI = 1.12

            if ui.toolBox.currentIndex() != 4 or (ui.toolBox.currentIndex() == 4 and bvp == bnp and tvp == tnp):
                text_abzac('Ж.3    Значение коэффициента ψ следует вычислять по формулам таблиц Ж.1 и Ж.2 в зависимости от числа закреплений сжатого пояса, вида нагрузки и места ее приложения, а также от коэффициента α, равного:')
                # text_abzac('для прокатных двутавров:')


                text_centr('α = 1.54 ∙ It / Iy ∙ (lef / h)^2')
                if ui.toolBox.currentIndex() == 0:
                    It = round(kI/3 * (2 * element['b'] * (element['t']**3) + (element['h'] - 2 * element['t']) * (element['s']**3)))
                if ui.toolBox.currentIndex() == 1:
                    if s <= D / 10:
                        It = pi * dsr**3 * s / 4
                    else:
                        It = (pi * D**4 / 32) * (1 - (d / D)**4)
                if ui.toolBox.currentIndex() == 2 or ui.toolBox.currentIndex() == 3:
                    It = round(kI/3 * ((2 * h * tw**3) + (2 * b * tw**3)),)
                if ui.toolBox.currentIndex() == 4:
                    if bvp == bnp and tvp == tnp:
                        It = round(kI/3 * ((hw * tw**3) + (bvp * tvp**3) + (bnp * tnp**3)))


                lef = round(sum(Lbk) * 1000)
                if ui.toolBox.currentIndex() != 4:
                    alfa = round(1.54 * It / Iy * (lef / h)**2, 4)
                if ui.toolBox.currentIndex() == 4:
                    if bvp == bnp and tvp == tnp:
                        alfa = round(8 * ((lef * tvp /(h * bvp))**2 * (1+((0.5 * h * tw**3) / (bvp * tvp**3)))), 4)

                print('alfa = ', alfa)

                text_abzac('где:')
                if ui.toolBox.currentIndex() == 0:
                    text_abzac('It = k / 3 ∙ (2 ∙ b ∙ t^3 + (h - 2 ∙ t) ∙ s^3) = {} мм4  -  момент инерции при свободном кручении:'.format(It))
                if ui.toolBox.currentIndex() == 1:
                    if s <= D / 10: 
                        text_abzac('It = 3.1416 ∙ Dsr^3 ∙ s / 4 = {} мм4 - момент инерции при свободном кручении:'.format(It))
                    else:
                        text_abzac('It = (3.1416 ∙ Dsr^4 / 32) ∙ (1 - (d / D)^4) = {} мм4 - момент инерции при свободном кручении:'.format(It))
                if ui.toolBox.currentIndex() == 2 or ui.toolBox.currentIndex() == 3:
                    text_abzac('It = k / 3 ∙ ((2 ∙ h ∙ t^3) + (2 ∙ b ∙ t^3)) = {} мм4  -  момент инерции при свободном кручении:'.format(It))
                if ui.toolBox.currentIndex() == 4:
                    if bvp == bnp and tvp == tnp:
                        text_abzac('It = k / 3 ∙ ((hw ∙ tw^3) + (bvp ∙ tvp^3) + (bnp ∙ tnp^3)) = {} мм4  -  момент инерции при свободном кручении:'.format(It))

                text_abzac('k = 1.29 (1.12)  -  для двутавра (для швеллера);')

                if ui.toolBox.currentIndex() == 0:
                    text_abzac('h = {} мм, b = {} мм, t = {} мм, s = {} мм  -  размеры профиля по ГОСТу.'.format(element['h'], element['b'], element['t'], element['s']))
                if ui.toolBox.currentIndex() == 1:
                    text_abzac('D = {} мм, d = {} мм, s = {} мм  -  размеры трубы (наружный и внутренний диаметр, толщина стенки).'.format(element['h'], element['b'], element['t'], element['s']))
                if ui.toolBox.currentIndex() == 2 or ui.toolBox.currentIndex() == 3:
                    text_abzac('h = {} мм, b = {} мм, t = {} мм  -  размеры профиля по ГОСТу.'.format(h, b, tw))
                if ui.toolBox.currentIndex() == 4:
                    if bvp == bnp and tvp == tnp:
                        text_abzac('hw = {} мм, tw = {} мм, bvp = {} мм, tvp = {} мм, bnp = {} мм, tnp = {} мм  -  размеры профиля.'.format(hw, tw, bvp, tvp, bnp, tnp))

                text_abzac('Iy = {} мм4  -  момент инерции'.format(Iy))
                text_abzac('lef = {} мм  -  расстояние между точками закреплений сжатого пояса от поперечных смещений.'.format(lef))
                text_centr('α = 1.54 ∙ {} мм4 / {} мм4 ∙ ({} мм / {} мм)^2 = {}'.format(It, Iy, lef, h, alfa))

                if alfa > 400:
                    text_centr_red('Т. к.  α > 400  -  необходимо увеличить сечение')
                    print('alfa > 400', alfa)
                    return
                if alfa < 0.1:
                    text_centr_red('Т. к.  α < 0.1  -  необходимо увеличить сечение')
                    print('alfa < 0.1', alfa)
                    return

                nZakrepName = ui.comboBox_4.currentIndex()
                if nZakrepName == 0: nZakrep = 'без закреплений'
                if nZakrepName == 1: nZakrep = 'два и более, делящие пролет балки на равные части'
                if nZakrepName == 2: nZakrep = 'одно в середине'

                VidNagrName = ui.comboBox_5.currentIndex()

                ui.textEdit_3.append('')
                text_abzac('Число закреплений сжатого пояса в пролете - ' + str(nZakrep) + '.')
                ui.textEdit_3.insertPlainText('.')

                if nZakrep == 'без закреплений':
                    if VidNagrName == 0: VidNagr = 'равномерно распределенная'
                    if VidNagrName == 1: VidNagr = 'сосредоточенная'
                    text_abzac('Вид нагрузки в пролете - ' + str(VidNagr) + '.')
                    text_abzac('Пояс, к которому приложена нагрузка - сжатый')
                    ui.textEdit_3.insertPlainText('.')

                    if VidNagr == 'сосредоточенная':
                        if 0.1 <= alfa <= 40:
                            trizybec = round(1.75 + 0.09 * alfa, 4)
                            text_centr('ψ = 1.75 + 0.09 ∙ α = 1.75 + 0.09 ∙ {} = {}'.format(alfa, trizybec))
                        if 40 < alfa <= 400:
                            trizybec = round(3.3 + 0.053 * alfa - 4.5*10**-5 * alfa**2, 4)
                            text_centr('ψ = 3.3 + 0.053 ∙ α - 4.5 ∙ 10^-5 ∙ α^2 = 3.3 + 0.053 ∙ {} - 4.5 ∙ 10^-5 ∙ {}^2 = {}'.format(alfa, alfa, trizybec))
                    if VidNagr == 'равномерно распределенная':
                        if 0.1 <= alfa <= 40:
                            trizybec = round(1.6 + 0.08 * alfa, 4)
                            text_centr('ψ = 1.6 + 0.08 ∙ α = 1.6 + 0.08 ∙ {} = {}'.format(alfa, trizybec))
                        if 40 < alfa <= 400:
                            trizybec = round(3.15 + 0.04 * alfa - 2.7*10**-5 * alfa**2, 4)
                            text_centr('ψ = 3.15 + 0.04 ∙ α - 2.7 ∙ 10^-5 ∙ α^2 = 3.15 + 0.04 ∙ {} - 2.7 ∙ 10^-5 ∙ {}^2 = {}'.format(alfa, alfa, trizybec))

                if nZakrep == 'два и более, делящие пролет балки на равные части':
                    text_abzac('Пояс, к которому приложена нагрузка - любой')
                    # text_abzac_insert_color('любой')
                    text_abzac('Пояс, к которому приложена нагрузка - сжатый')
                    # text_abzac_insert_color('сжатый')
                    ui.textEdit_3.insertPlainText('.')
                    if 0.1 <= alfa <= 40:
                        trizybec = round(2.25 + 0.07 * alfa, 4)
                        text_centr('ψ = 2.25 + 0.07 ∙ α = 2.25 + 0.07 ∙ {} = {}'.format(alfa, trizybec))
                    if 40 < alfa <= 400:
                        trizybec = round(3.6 + 0.04 * alfa - 3.5*10**-5 * alfa**2, 4)
                        text_centr('ψ = 3.6 + 0.04 ∙ α - 3.5 ∙ 10^-5 ∙ α^2 = 3.6 + 0.04 ∙ {} - 3.5 ∙ 10^-5 ∙ {}^2 = {}'.format(alfa, alfa, trizybec))

                if nZakrep == 'одно в середине':
                    if VidNagrName == 0: VidNagr = 'сосредоточенная в середине'
                    if VidNagrName == 1: VidNagr = 'сосредоточенная в четверти'
                    if VidNagrName == 2: VidNagr = 'равномерно распределенная'

                    text_abzac('Вид нагрузки в пролете - ' + str(VidNagr) + '.')
                    text_abzac('Пояс, к которому приложена нагрузка - сжатый')

                    if VidNagr == 'сосредоточенная в середине':
                        ka = 1.75
                        if 0.1 <= alfa <= 40:
                            trizybec = round(ka * (2.25 + 0.07 * alfa), 4)
                            text_centr('ψ = {} ∙ (2.25 + 0.07 ∙ α) = {} ∙ (2.25 + 0.07 ∙ {}) = {}'.format(ka, ka, alfa, trizybec))
                        if 40 < alfa <= 400:
                            trizybec = round(ka * (3.6 + 0.04 * alfa - 3.5*10**-5 * alfa**2), 4)
                            text_centr('ψ = {} ∙ (3.6 + 0.04 ∙ α - 3.5 ∙ 10^-5 ∙ α^2) = {} ∙ (3.6 + 0.04 ∙ {} - 3.5 ∙ 10^-5 ∙ {}^2) = {}'.format(ka, ka, alfa, alfa, trizybec))
                    if VidNagr == 'сосредоточенная в четверти' or VidNagr == 'равномерно распределенная':
                        ka = 1.14
                        if 0.1 <= alfa <= 40:
                            trizybec = round(ka * (2.25 + 0.07 * alfa), 4)
                            text_centr('ψ = {} ∙ (2.25 + 0.07 ∙ α) = {} ∙ (2.25 + 0.07 ∙ {}) = {}'.format(ka, ka, alfa, trizybec))
                        if 40 < alfa <= 400:
                            trizybec = round(ka * (3.6 + 0.04 * alfa - 3.5*10**-5 * alfa**2), 4)
                            text_centr('ψ = {} ∙ (3.6 + 0.04 ∙ α - 3.5 ∙ 10^-5 ∙ α^2) = {} ∙ (3.6 + 0.04 ∙ {} - 3.5 ∙ 10^-5 ∙ {}^2) = {}'.format(ka, ka, alfa, alfa, trizybec))

                ui.textEdit_3.append('')
                text_abzac('Ж.2    Коэффициент φb для расчета на устойчивость по 8.4.1, определяется через значение φ1:')
                text_centr('φ1 =  ψ ∙ Iy / Ix ∙ (hy / lef)^2 ∙ E / Ry')

                fi1 = round(trizybec * Iy / Ix * (h / lef)**2 * E / Ry, 4)
                print('fi1 = ', fi1)


                text_abzac('где:')
                text_abzac('ψ = {}  -  коэффициент, вычисляемый согласно требованиям Ж.3;'.format(trizybec))
                text_abzac('Iy = {} мм4, Ix = {} мм4  -  момент инерции сечения балки'.format(Iy, Ix))

                if ui.toolBox.currentIndex() == 0:
                    text_abzac('hy = {} мм  -  полная высота сечения балки.'.format(element['h']))
                if ui.toolBox.currentIndex() == 2 or ui.toolBox.currentIndex() == 3 or ui.toolBox.currentIndex() == 4:
                    text_abzac('hy = {} мм  -  полная высота сечения балки.'.format(h))

                text_abzac('lef = {} мм  -  расстояние между точками закреплений сжатого пояса от поперечных смещений;'.format(lef))
                text_abzac('E = {} Н/мм2  -  модуль упругости стали;'.format(E))
                text_abzac('Ry = {} Н/мм2  -  расчетное сопротивление растяжению, сжатию, изгибу по пределу текучести.'.format(Ry))
                text_centr('φ1 =  {} ∙ {} мм4 / {} мм4 ∙ ({} мм / {} мм)^2 ∙ {} Н/мм2 / {} Н/мм2 = {}'.format(trizybec, Iy, Ix , h, lef, E, Ry, fi1))

                if TipSech == ProfALL[1] or TipSech == ProfALL[2] and ui.toolBox.currentIndex() == 0: # если швеллер
                    ui.textEdit_3.append('')
                    text_abzac('Ж.7    Для балки швеллерного сечения коэффициент φb следует принимать равным:')
                    fib = round(0.7 * fi1, 4)
                    text_centr(' φb = 0.7 ∙ φ1 =  0.7 ∙ {}  = {}'.format(fi1, fib))

                # if TipSech == ProfALL[0] and ui.toolBox.currentIndex() != 1: # если двутавр
                if TipSech == ProfALL[0] and ui.toolBox.currentIndex() == 0: # если двутавр
                    ui.textEdit_3.append('')
                    text_abzac('Для балки и консоли коэффициент φb следует принимать равным:')
                    if fi1 <= 0.85:
                        fib = fi1
                        text_centr('Т.к. φ1 ≤ 0.85:  φb = φ1 = {}'.format(fib))
                    else:
                        fib = round(0.68 + 0.21 * fi1, 4)
                        text_centr('Т.к. φ1 > 0.85:  φb = 0.68 + 0.21 ∙ φ1 = 0.68 + 0.21 ∙ {} = {}'.format(fi1, fib))
                    print('fib = ', fib)


                if ui.toolBox.currentIndex() == 2 or ui.toolBox.currentIndex() == 3 or ui.toolBox.currentIndex() == 4:
                    ui.textEdit_3.append('')
                    text_abzac('Для балки и консоли коэффициент φb следует принимать равным:')
                    if fi1 <= 0.85:
                        fib = fi1
                        text_centr('Т.к. φ1 ≤ 0.85:  φb = φ1 = {}'.format(fib))
                    else:
                        fib = round(0.68 + 0.21 * fi1, 4)
                        text_centr('Т.к. φ1 > 0.85:  φb = 0.68 + 0.21 ∙ φ1 = 0.68 + 0.21 ∙ {} = {}'.format(fi1, fib))
                    print('fib      ======== = ', fib)


            if ui.toolBox.currentIndex() == 4:
                if bvp != bnp or tvp != tnp:
                    if Avp >= Anp:
                        h1 = round(abs(Yco_vp), 2)
                        h2 = round(abs(Yco_np), 2)
                        I1 = Ix_vp
                        I2 = Ix_np
                        if ui.comboBox_5.currentText() == ChisloZakrepSD[0]:
                            B = round(sigma - 1, 2)
                    else:
                        h1 = round(abs(Yco_np), 2)
                        h2 = round(abs(Yco_vp), 2)
                        I1 = Ix_np
                        I2 = Ix_vp
                    It = round(1.29/3 * ((hw * tw**3) + (bvp * tvp**3) + (bnp * tnp**3)))
                    lef = round(sum(Lbk) * 1000)
                    alfa = round(1.54 * It / Iy * (lef / h)**2, 4)
                    b1 = bvp
                    b2 = bnp
                    hs = h
                    h = h2w
                    # n = round(Ix / (I1 + I2), 3)
                    n = round(I1 / (I1 + I2), 3)
                    betta = round((2 * n - 1) * (0.47 - 0.035 * (b1 / h) * (1 + (b1 / h) - 0.072 * (b1 / h)**2)), 3)
                    eta = round((1 - n) * (9.87 * n + 0.385 * (It / I2) * (lef / h)**2), 3)
                    sigma = round(n + 0.734 * betta, 3)
                    nuy = round(n + 1.145 * betta, 3)
                    
                    '''Приложение нагрузки на верхниий пояс'''
                    if Avp >= Anp:      # Увеличенный верхний пояс
                        '''при нагрузке сосредоточенной в середине пролета'''
                        if ui.comboBox_5.currentText() == VidNagrSD[0]:
                            B = round(sigma - 1, 3)
                        '''при нагрузке равномерно распределенной'''
                        if ui.comboBox_5.currentText() == VidNagrSD[1]:
                            B = round(nuy - 1, 3)
                        '''при нагрузке вызывающей чистый изгиб '''
                        if ui.comboBox_5.currentText() == VidNagrSD[2]:
                            B = round(betta, 3)
                    else:      # уменьшенный верхний пояс
                        if ui.comboBox_5.currentText() == VidNagrSD[0]:
                            B = round(- sigma, 3)
                        '''при нагрузке равномерно распределенной'''
                        if ui.comboBox_5.currentText() == VidNagrSD[1]:
                            B = (- nuy, 3)
                        '''при нагрузке вызывающей чистый изгиб '''
                        if ui.comboBox_5.currentText() == VidNagrSD[2]:
                            B = round(- betta, 3)

                    if n <= 0.9:
                        if ui.comboBox_5.currentText() == VidNagrSD[0]:
                            C = round(0.330 * eta, 3)
                            D = 3.265
                        if ui.comboBox_5.currentText() == VidNagrSD[1]:
                            C = round(0.481 * eta, 3)
                            D = 2.247
                        if ui.comboBox_5.currentText() == VidNagrSD[2]:
                            C = round(0.101 * eta, 3)
                            D = 4.315
                    if n > 0.9:
                        if ui.comboBox_5.currentText() == VidNagrSD[0]:
                            C1 = 0.330 * eta
                            C2 = 0.826 * alfa
                            D = 3.265
                        if ui.comboBox_5.currentText() == VidNagrSD[1]:
                            C1 = 0.481 * eta
                            C2 = 0.1202 * alfa
                            D = 2.247
                        if ui.comboBox_5.currentText() == VidNagrSD[2]:
                            C1 = 0.101 * eta
                            C2 = 0.0253 * alfa
                            D = 4.315
                        C = round(interpoi([0.9, 1.0], [C1, C2], n), 3)

                    tria = round((B + (B**2 + C)**0.5) * D, 3)

                    ui.textEdit_3.append('')
                    text_abzac('Ж.5    Значения В, С и D в формуле (Ж.9) определяют по таблицам Ж.4 и Ж.5 в зависимости от коэффициентов:')
                    text_centr('α = 1.54 ∙ It / Iy ∙ (lef / hs)^2')
                    text_centr('n =  I1 / I1 + I2')
                    text_centr('β = (2 ∙ n - 1) ∙ (0.47 - 0.035 ∙ (b1 / h) ∙ (1 + (b1 / h) - 0.072 ∙ (b1 / h)^2))')
                    text_centr('η = (1 - n) ∙ (9.87 ∙ n + 0.385 ∙ (It / I2) ∙ (lef / h)^2)')
                    text_centr('δ = n + 0.734 ∙ β')
                    text_centr('µ = n + 1.145 ∙ β')

                    text_abzac('где:')
                    text_abzac('It = k / 3 ∙ ((hw ∙ tw^3) + (bvp ∙ tvp^3) + (bnp ∙ tnp^3)) = {} мм4  -  момент инерции при свободном кручении:'.format(It))
                    text_abzac('Iy = {} мм4  -  момент инерции'.format(Iy))
                    text_abzac('hs = {} мм  -  полная высота сечения;'.format(hs))
                    text_abzac('lef = {} мм  -  расстояние между точками закреплений сжатого пояса от поперечных смещений;'.format(lef))
                    text_abzac('hw = {} мм, tw = {} мм, bvp = {} мм, tvp = {} мм, bnp = {} мм, tnp = {} мм  -  размеры профиля.'.format(hw, tw, bvp, tvp, bnp, tnp))
                    text_abzac('I1 = {} мм4 и I2 = {} мм4 -  моменты инерции сечения более развитого и менее развитого поясов относительно оси симметрии сечения балки соответственно;'.format(I1, I2))
                    text_abzac('b1 = {} мм  -  ширина верхнего поясов;'.format(bvp))
                    text_abzac('h = {} мм  -  расстояние между осями поясов;'.format(h2w))

                    text_centr('α = 1.54 ∙ {} мм4 / {} мм4 ∙ ({} мм / {} мм)^2 = {}'.format(It, Iy, lef, hs, alfa))
                    text_centr('n =  {} мм4 / {} мм4 + {} мм4 = {}'.format(I1, I1, I2, n))
                    text_centr('β = (2 ∙ {} - 1) ∙ (0.47 - 0.035 ∙ ({} мм / {} мм) ∙ (1 + ({} мм / {} мм) - 0.072 ∙ ({} мм / {} мм)^2)) = {}'.format(n, b1, h, b1, h, b1, h, betta))
                    text_centr('η = (1 - {}) ∙ (9.87 ∙ {} + 0.385 ∙ ({} мм4 / {} мм4) ∙ ({} мм / {} мм)^2) = {}'.format(n, n, It, I2, lef, h, eta))
                    text_centr('δ = {} + 0.734 ∙ {} = {}'.format(n, betta, sigma))
                    text_centr('µ = {} + 1.145 ∙ {} = {}'.format(n, betta, nuy))


                    fi1 = tria * Iy / Ix * (2 * h * h1 / (lef)**2) * E / Ry
                    fi2 = tria * Iy / Ix * (2 * h * h2 / (lef)**2) * E / Ry
                    fi1 = round(fi1, 2)
                    fi2 = round(fi2, 2)




                    ui.textEdit_3.append('')
                    text_abzac('Ж.4    Для разрезной балки двутаврового сечения с одной осью симметрии φb следует оределять с помощью коэффициентов φ1, φ2 и n:')
                    text_centr('φ1 =  ψa ∙ Iy / Ix ∙ (2 ∙ h ∙ h1 / (lef)^2) ∙ E / Ry')
                    text_centr('φ2 =  ψa ∙ Iy / Ix ∙ (2 ∙ h ∙ h2 / (lef)^2) ∙ E / Ry')

                    text_abzac('где:')
                    text_abzac('В = {}, С = {} и D = {}  -  полученные значения в зависимости от коэффициентов по Ж.5;'.format(B, C, D))
                    text_abzac('ψa =  (B + (B^2 + C)^0.5) ∙ D = ({} + ({}^2 + {})^0.5) ∙ {} = {}  -  коэффициент;'.format(B, B, C, D, tria))
                    text_abzac('Ix = {} мм4 и Iy = {} мм4  -  моменты инерции сечения балки;'.format(Ix, Iy))
                    text_abzac('h = {} мм  -  расстояние между осями поясов;'.format(h2w))
                    text_abzac('h1 = {} мм и h2 = {} мм  -  расстояние от центра тяжести сечения до оси более и менее развитого поясов;'.format(h1, h2))
                    text_abzac('lef = {} мм  -  расстояние между точками закреплений сжатого пояса от поперечных смещений;'.format(lef))
                    text_abzac('E = {} Н/мм2  -  модуль упругости стали;'.format(E))
                    text_abzac('Ry = {} Н/мм2  -  расчетное сопротивление стали.'.format(Ry))

                    text_centr('φ1 =  {} ∙ {} мм4 / {} мм4 ∙ (2 ∙ {} мм ∙ {} мм / ({} мм)^2) ∙ {} Н/мм2 / {} Н/мм2 = {}'.format(tria, Iy, Ix, h, h1, lef, E, Ry, fi1))
                    text_centr('φ2 =  {} ∙ {} мм4 / {} мм4 ∙ (2 ∙ {} мм ∙ {} мм / ({} мм)^2) ∙ {} Н/мм2 / {} Н/мм2 = {}'.format(tria, Iy, Ix, h, h2, lef, E, Ry, fi2))


                    if Avp < Anp and 5 <= lef / b2 <= 25:
                        yminsh = 1.025 - 0.015 * lef / b2
                        if yminsh > 0.95: 
                            yminsh = 0.95
                        fi2 = fi2 * yminsh
                        fi2 = round(fi2, 2)

                    '''Если сжатый пояс (верхний) более развитый'''
                    if Avp < Anp and lef / b2 > 25:
                        text_centr_red('Значение lef / b2 > 25 в балках с менее развитым верхним поясом не допустимо')
                        return

                    '''Если сжатый пояс более развитый'''
                    if Avp >= Anp and fi2 <= 0.85:
                        fib = fi1
                        fib = round(fib, 2)
                        if fi1 > 1:
                            text_centr_red('Т.к. по таблице Ж.3 φ2 ≤ 0.85:  φb = φ1 = {} > 1  -  условие НЕ выполено, необходимо изменить сечение балки'.format(fi1))
                            return
                        else:
                            text_centr('Т.к. по таблице Ж.3 φ2 ≤ 0.85:  φb = φ1 = {} ≤ 1'.format(fib))
                    
                    if Avp >= Anp and fi2 > 0.85:
                        fib = fi1 * (0.21 + 0.68 * ((n / fi1) + ((1 - n) / fi2)))
                        fib = round(fib, 2)
                        text_centr('Т.к. по таблице Ж.3 φ2 > 0.85:  φb = φ1 ∙ (0.21 + 0.68 ∙ ((n / φ1) + ((1 - n) / φ2))) ≤ 1')
                        if fib > 1:
                            text_centr_red('φb = {} ∙ (0.21 + 0.68 ∙ (({} / {}) + ((1 - {}) / {}))) = {} > 1  -  условие НЕ выполено,'.format(fi1, n, fi1, n, fi2, fib))
                            return
                        else:
                            text_centr('φb = {} ∙ (0.21 + 0.68 ∙ (({} / {}) + ((1 - {}) / {}))) = {} ≤ 1'.format(fi1, n, fi1, n, fi2, fib))

                    '''Если сжатый пояс менее развитый'''
                    if Avp < Anp and fi2 <= 0.85:
                        fib = fi2
                        fib = round(fib, 2)
                        text_centr('Т.к. по таблице Ж.3 φ2 ≤ 0.85:  φb = φ2 = {}'.format(fib))

                    if Avp < Anp and fi2 > 0.85:
                        fib = 0.68 + 0.21 * fi2
                        fib = round(fib, 2)
                        text_centr('Т.к. по таблице Ж.3 φ2 > 0.85:  φb = 0.68 + 0.21 ∙ φ2} ≤ 1')
                        if fib > 1:
                            text_centr_red('φb = 0.68 + 0.21 ∙ {} = {} > 1  -  условие НЕ выполено, откорректируйте сечение балки'.format(fi2, fib))
                            return
                        else:
                            text_centr('φb = 0.68 + 0.21 ∙ {} = {} ≤ 1'.format(fi2, fib))




            ui.textEdit_3.append('')
            text_abzac('8.4.1    Расчет на устойчивость балок 1-го класса, а также бистальных балок 2-го класса, следует выполнять по формуле:')
            text_centr('Mx / (φb ∙ Wcx ∙ Ry ∙ yc) ≤ 1')

            Ystouchnvost = round(MmaxABS / (fib * Wx * Ry * yc), 4)

            text_abzac('где:')
            text_abzac('Mx = {} H мм  -  абсолютное значения изгибающего момента;'.format(MmaxABS))
            text_abzac('φb = {}  -  коэффициент устойчивости при изгибе, определяемый по приложению Ж для балок с опорными сечениями, закрепленными от боковых смещений и поворота;'.format(fib))
            text_abzac('Wcx = Wx = {} мм3  -  момент сопротивления сечения относительно оси х-х, вычисленный для наиболее сжатого волокна сжатого пояса;'.format(Wx))
            text_abzac('Ry = {} Н/мм2  -  расчетное сопротивление стали сжатого пояса;'.format(Ry))
            text_abzac('yc = {}  -  коэффициент условий работы.'.format(yc))
            if Ystouchnvost <= 1:
                text_centr('{} H мм / ({} ∙ {} мм3 ∙ {} Н/мм2 ∙ {}) = {} ≤ 1  -  условие выполняется'.format(MmaxABS, fib, Wx, Ry, yc, Ystouchnvost))
                print('Ystouchnvost <= 1 = ', Ystouchnvost)
            else:
                text_centr_red('{} H мм / ({} ∙ {} мм3 ∙ {} Н/мм2 ∙ {}) = {} ≥ 1  -  условие НЕ выполняется'.format(MmaxABS, fib, Wx, Ry, yc, Ystouchnvost))
                print('Ystouchnvost > 1 = ', Ystouchnvost)
                return

            ui.textEdit_3.append('')

    # if ui.toolBox.currentIndex() != 1:
    text_abzac('При расчете по второй группе предельных состояний определяем предельные и расчетные прогибы балки.')
    ui.textEdit_3.append('')
    text_abzac('Расчетный прогиб балки между опорами:')
    
    EJvIIRabsH = EJvIIRabs * 9806.65 * 10**9 #(Н мм3)
    RfR = round(EJvIIRabsH / (E * Ix), 2)
    print('RfR = ', RfR)
    

    text_centr('f = EIf / (E ∙ I)')
    text_abzac('где:')
    text_abzac('EIf = {} тс м3 = {} Н мм3  -  прогиб балки между опорами, согласно эпюре прогибов EIf;'.format(EJvIIRabs, EJvIIRabsH))
    text_abzac('E = {} Н/мм2  -  модуль упругости стали;'.format(E))
    text_abzac('I = {} мм4  -  момент инерции сечения балки.'.format(Ix))
    text_centr('f = {} Н мм3 / ({} Н/мм2 ∙ {} мм4) = {} мм'.format(EJvIIRabsH, E, Ix, RfR))
    ui.textEdit_3.append('')

    if sum(Lkl) != 0:
        text_abzac('Расчетный прогиб балки на левом консольном участке:')

        EJvIIK1absH = EJvIIK1abs * 9806.65 * 10**9 #(Н мм3)
        fRR = round(EJvIIK1absH / (E * Ix), 2)
        print('fRR = ', fRR)
        text_centr('fo = EIf_kl / (E ∙ I)')
        text_abzac('где:')
        text_abzac('EIf_kl = {} тс м3 = {} Н мм3  -  прогиб балки левой консоли, согласно эпюре прогибов EIf;'.format(EJvIIK1abs, EJvIIK1absH))
        text_abzac('E = {} Н/мм2  -  модуль упругости стали;'.format(E))
        text_abzac('I = {} мм4  -  момент инерции сечения балки.'.format(Ix))
        text_centr('fo = {} Н мм3 / ({} Н/мм2 ∙ {} мм4) = {} мм'.format(EJvIIK1absH, E, Ix, fRR))
        ui.textEdit_3.append('')

    if sum(Lkp) != 0:
        text_abzac('Расчетный прогиб балки на правом консольном участке:')

        EJvIIK2absH = EJvIIK2abs * 9806.65 * 10**9 #(Н мм3)
        RRf = round(EJvIIK2absH / (E * Ix), 2)
        print('RRf = ', RRf)

        text_centr('fp = EIf_kp / (E ∙ I)')
        text_abzac('где:')
        text_abzac('EIf_kp = {} тс м3 = {} Н мм3  -  прогиб балки правой консоли, согласно эпюре прогибов EIf;'.format(EJvIIK2abs, EJvIIK2absH))
        text_abzac('E = {} Н/мм2  -  модуль упругости стали;'.format(E))
        text_abzac('I = {} мм4  -  момент инерции сечения балки.'.format(Ix))
        text_centr('fp = {} Н мм3 / ({} Н/мм2 ∙ {} мм4) = {} мм'.format(EJvIIK2absH, E, Ix, RRf))
        ui.textEdit_3.append('')

    kfu = dannie[0]
    Lbk = sum(Lbk) * 1000
    fu = round(Lbk / kfu, 2)

    text_abzac('Предельный прогиб балки между опорами балки')
    text_centr('fu = L / {} = {} мм / {} = {} мм'.format(kfu, Lbk, kfu, fu))
    if sum(Lkl) != 0:
        Lklabs = sum(Lkl) * 1000
        ful = round(Lklabs / 150, 2)
        ui.textEdit_3.append('')
        text_abzac('Предельный прогиб балки на левом консольном участке')
        text_centr('fuo = Lo / {} = {} мм / {} = {} мм'.format(150, Lklabs, 150, ful))
    if sum(Lkp) != 0:
        Lkpabs = sum(Lkp) * 1000
        fup = round(Lkpabs / 150, 2)
        ui.textEdit_3.append('')
        text_abzac('Предельный прогиб балки на правом консольном участке')
        text_centr('fup = Lp / {} = {} мм / {} = {} мм'.format(150, Lkpabs, 150, fup))

    text_abzac('где:')
    text_abzac('L = {} мм  -  расстояние между опорами балки'.format(Lbk))
    if Lkl != 0:
        text_abzac('Lo = {} мм  -  расстояние левой консоли балки'.format(sum(Lkl)))
    if Lkp != 0:
        text_abzac('Lp = {} мм  -  расстояние правой консоли балки'.format(sum(Lkp)))
    ui.textEdit_3.append('')

    # text_abzac('15.1.1    При расчете строительных конструкций должно быть выполнено условие по СП 20.13330.2016  -  f ≤ fu, где: f и fu  -  расчетный и предельный прогиб балки:')
    text_abzac('15.1.1    При расчете строительных конструкций должно быть выполнено условие по СП 20.13330.2016:')
    
    if RfR <= fu:
        text_centr('f = {} мм ≤ fu = {} мм  -  условие выполнено'.format(RfR, fu))
    else:
        text_centr_red('f = {} мм > fu = {} мм  -  условие НЕ выполнено'.format(RfR, fu))
        return
    
    if sum(Lkl) != 0:
        if fRR <= ful:
            text_centr('fo = {} мм ≤ fuo = {} мм  -  условие выполнено'.format(fRR, ful))
        else:
            text_centr_red('fo = {} мм > fuo = {} мм  -  условие НЕ выполнено'.format(fRR, ful))
            return
    
    if sum(Lkp) != 0:
        if RRf <= fup:
            text_centr('fp = {} мм ≤ fup = {} мм  -  условие выполнено'.format(RRf, fup))
        else:
            text_centr_red('fp = {} мм > fup = {} мм  -  условие НЕ выполнено'.format(RRf, fup))
            return
    
    # ///////////////////////////////////////////////////////////////////////////////////////////////

def podbor():
    global G, Ystouchnvost, RfR, fu, fRR, ful, RRf, fup, Tay, alfa, Stop841

    def WinMassage(TextMessage):
        QtWidgets.QMessageBox.information(Form, 'Подобр сечения', TextMessage)

    if ui.toolBox.currentIndex() == 1:
        raschet()
        return WinMassage('Выполен расчет трубы по заданным размерам')
    if ui.toolBox.currentIndex() == 4:
        raschet()
        return WinMassage('Выполен расчет составного двутавра по заданным размерам')

    print('////nnnn////  ', ui.comboBox_2.currentIndex())

    raschet()

    def perebor(W=1):
        if ui.toolBox.currentIndex() == 0:
            ui.comboBox_2.setCurrentIndex(ui.comboBox_2.currentIndex() + W)
            print('Сечение №', ui.comboBox_2.currentText().strip(), '  -  ', ui.comboBox_2.currentIndex())
        if ui.toolBox.currentIndex() == 2:
            ui.comboBox_6.setCurrentIndex(ui.comboBox_6.currentIndex() + W)
            print('Сечение №', ui.comboBox_6.currentText().strip(), '  -  ', ui.comboBox_6.currentIndex())
        if ui.toolBox.currentIndex() == 3:
            ui.comboBox_7.setCurrentIndex(ui.comboBox_7.currentIndex() + W)
            print('Сечение №', ui.comboBox_7.currentText().strip(), '  -  ', ui.comboBox_7.currentIndex())
        raschet()

    try:
        if G == 0:
            error_show('Укажите нагрузку на балку')
            return
    except:
        error_show('Укажите длину балки')
        return

    def ysloviePlus():
        while G > 1:
            perebor()
        while Tay > 1:
            perebor()
        
        if ui.toolBox.currentIndex() != 1 or Stop841 == 'No':
            while alfa < 0.1 or alfa > 400:
                perebor()
            while Ystouchnvost > 1:
                perebor()
        
        while RfR > fu:
            perebor()
        if sum(Lkl) != 0:
            while fRR > ful:
                perebor()
        if sum(Lkp) != 0:
            while RRf > fup:
                perebor()

    ysloviePlus()
    
    '''Остановка перебора, если доходим до низа списка сечения в сортаменте'''
    def stoplist():
        if ui.toolBox.currentIndex() == 0 and ui.comboBox_2.currentIndex() == 0:
            eval('break')
        if ui.toolBox.currentIndex() == 2 and ui.comboBox_6.currentIndex() == 0:
            eval('break')
        if ui.toolBox.currentIndex() == 3 and ui.comboBox_7.currentIndex() == 0:
            eval('break')

    try:
        while RfR < fu:
            stoplist()
            del RfR
            perebor(-1)
    except: ysloviePlus()
    ysloviePlus()


    if sum(Lkl) != 0:
        try:
            while fRR < ful:
                stoplist()
                del fRR
                perebor(-1)
        except: ysloviePlus()
        ysloviePlus()
        
    if sum(Lkp) != 0:
        try:
            while RRf < fup:
                stoplist()
                del RRf
                perebor(-1)
        except: ysloviePlus()
        ysloviePlus()

    try: 
        del G, Ystouchnvost, RfR, fu, fRR, ful, RRf, fup, Tay, alfa, Stop841
    except: pass

    if ui.toolBox.currentIndex() == 0:
        TextMessage = 'Подобрано сечение № {} - {}'.format(ui.comboBox_2.currentText().strip(), ui.comboBox_1.currentText().strip())
    if ui.toolBox.currentIndex() == 2:
        TextMessage = 'Подобрано сечение № {} - {}'.format(ui.comboBox_6.currentText().strip(), 'Замкнутое сечение по ГОСТ 30245')
    if ui.toolBox.currentIndex() == 3:
        TextMessage = 'Подобрано сечение № {} - {}'.format(ui.comboBox_7.currentText().strip(), 'Замкнутое сечение по ГОСТ 30245')

    WinMassage(TextMessage)

'''# -----------------------------------------------------------------------------'''
'''# -----------------------------------------------------------------------------'''

def soxranka():
    # ff = QtWidgets.QFileDialog.getSaveFileName(Form, 'Сохранить как', QtCore.QDir.currentPath(), 'NMH ( *.nmh )')
    ff = QtWidgets.QFileDialog.getSaveFileName(Form, 'Сохранить как', QtCore.QDir.currentPath(), 'Balka ( *.blk )')
    katalog = ff [0]

    if katalog != '':       # при нажатии на кнопку ОТМЕНА, не происходит вылета проги
        sbor_dannih()
        # global Lkl, Lbk, Lkp, Nv, Nvx1, Nvx2, Moment, MomX, dannie, box_index
        savex = [Lkl, Lbk, Lkp, Nv, Nvx1, Nvx2, Moment, MomX, dannie, box_index]
        
        f = open(katalog, 'wb') # Запись в файл
        pickle.dump(savex, f) # помещаем объект в файл
        f.close()
        del savex # уничтожаем переменную savex  
        # ------------------------------------------------------------
        section = document.sections[0]
        section.top_margin  = Cm(2)     # верхнее поле
        section.bottom_margin = Cm(2)   # нижнее поле
        section.left_margin = Cm(3)     # левое поле
        section.right_margin = Cm(1.5)    # правое поле
        # ------------------------------------------------------------
        document.save('{}.docx'.format(katalog[:-4]))
        # Form.setWindowTitle(_translate("Form", "NMH - {}".format(katalog)))
        Form.setWindowTitle(_translate("Form", "Balka - {}".format(katalog)))
        ui.label_30.setText(_translate("Form", "Сохранение"))
        ui.textEdit_3.setText('')
        yvedomlenie('Данные сохранены в файл: {}'.format(katalog))
    else: return None

def otkrivalka():
    # ff = QtWidgets.QFileDialog.getOpenFileName(Form, 'Открыть', QtCore.QDir.currentPath(), 'NMH ( *.nmh )')
    ff = QtWidgets.QFileDialog.getOpenFileName(Form, 'Открыть', QtCore.QDir.currentPath(), 'Balka ( *.blk )')
    katalog = ff [0]
    if katalog != '':               # при нажатии на кнопку ОТМЕНА, не происходит вылета проги
        f = open(katalog, 'rb')     # чтение из файла
        loadx = pickle.load(f)      # извлекаем ообъект из файла
        
        for i in range(0, len(loadx)-1):
            for z in range(0, len(loadx[i])):
                if loadx[i][z] == 0:
                    loadx[i][z] = ''

        def vvod_open(tablica, stolbec, xList):
            _translate = QtCore.QCoreApplication.translate
            '''Очищаем все ячейки таблицы и узнаем количество строк её строк'''
            xxx = eval('[\'\' for z in range(0, ui.tableWidget{}.rowCount())]'.format(tablica))
            for i in range(0, len(xxx)):
                eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", str({})))'.format(tablica, i, stolbec, xxx[i]))
            '''Подставлеем значения из списков в каждую строку таблицы в выбранном столбце'''
            for i in range(0, len(xList)):
                eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", str({})))'.format(tablica, i, stolbec, xList[i]))
        
        vvod_open('_20', 0, loadx[0])
        vvod_open('_20', 1, loadx[1])
        vvod_open('_20', 2, loadx[2])
        
        vvod_open('_11', 0, loadx[3])
        vvod_open('_11', 1, loadx[4])
        vvod_open('_11', 2, loadx[5])
        
        vvod_open('_12', 0, loadx[6])
        vvod_open('_12', 1, loadx[7])

        vvod_open('_14', 0, loadx[8])
    
        '''
        Собираем боксы из индексов
        ui.comboBox_0 - НДС
        ui.comboBox_1 - Сечение балки профиль
        ui.comboBox_2 - Сечение балки номер
        ui.comboBox_3 - Марка стали
        ui.comboBox_4 - Число закреплений сжатого пояса в пролете
        ui.comboBox_5 - Вид нагрузки в пролете
        ui.comboBox_3_1 - Толщина стали
        '''

        ui.comboBox_0.setCurrentIndex(loadx[-1][0])

        ui.comboBox_1.setCurrentIndex(loadx[-1][1])
        ui.comboBox_2.clear()
        DYorSH()
        ui.comboBox_2.setCurrentIndex(loadx[-1][2])

        ui.comboBox_3.setCurrentIndex(loadx[-1][3])
        ui.comboBox_3_1.clear()
        TolStlist()
        ui.comboBox_3_1.setCurrentIndex(loadx[-1][6])

        ui.comboBox_4.setCurrentIndex(loadx[-1][4])
        ui.comboBox_5.clear()
        VidNaga()
        ui.comboBox_5.setCurrentIndex(loadx[-1][5])

        del loadx # уничтожаем переменную savex  
        ui.label_31.setText(_translate("Form", "Открытие"))
        ui.textEdit_3.setText('')
        yvedomlenie('Данные закгужены из файла: {}'.format(katalog))
        Form.setWindowTitle(_translate("Form", "NMH - {}".format(katalog)))
        return None
    else: return None

'''Авторизация'''
noyblock = 'no'
def block():
    global noyblock
    passeord = QtWidgets.QInputDialog.getText(Form, 'Авторизация', 'Введите пароль')
    # if passeord[0] == 'ckj;ysqgfhjkm':
    if passeord[0] == 'vxv':
        noyblock = 'Yes'
    else:
        noyblock = 'no'
        QtWidgets.QMessageBox.information(Form, 'Авторизация', 'Неверный пароль')

def raschet_block():
    global noyblock
    if noyblock == 'no':
        block()
    if noyblock == 'Yes':
        raschet()

def podbor_block():
    global noyblock
    if noyblock == 'no':
        block()
    if noyblock == 'Yes':
        podbor()

# -----------------------------------------------------------------------------

# ui.pushButton_10.clicked.connect (raschet_block)
# ui.pushButton_12.clicked.connect (podbor_block)
ui.pushButton_10.clicked.connect (raschet)
ui.pushButton_12.clicked.connect (podbor)
ui.pushButton_11.clicked.connect (soxranka)
ui.pushButton_8.clicked.connect (otkrivalka)
# ui.pushButton_9.clicked.connect (spravka)

# ui.pushButton_9.clicked.connect (testii)
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    sys.exit(app.exec_())
# -----------------------------------------------------------------------------